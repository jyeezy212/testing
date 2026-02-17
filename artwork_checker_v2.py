#!/usr/bin/env python3
"""
================================================================================
ARTWORK CHECKER v2.0.0 - COMPLETE REWRITE
================================================================================

A comprehensive tool for validating packaging artwork against copy documents.
Designed for use with ChatGPT Custom GPT for visual verification workflow.

FEATURES:
---------
✅ Copy document parsing with strikethrough detection
✅ Legacy text handling (` -> ` separator)
✅ PDF/AI text extraction via PyMuPDF
✅ Character-by-character comparison
✅ Smart zoom triggers for visual verification
✅ Comprehensive exclusion patterns for non-product text
✅ INCI capitalization validation
✅ Barcode scanning and validation
✅ Full-page annotated snapshots
✅ Markdown report generation
✅ Optional PDF export

USAGE:
------
    python artwork_checker_v2.py --copy <copy.docx> --artwork <artwork.pdf>
    
    ChatGPT Workflow:
    1. Execute script to generate automated report
    2. Review flagged fields requiring visual verification
    3. Visually verify and update results
    4. Present final corrected report

AUTHOR: Built for Amika brand artwork verification
VERSION: 2.0.0
DATE: 2025
================================================================================
"""

# =============================================================================
# IMPORTS
# =============================================================================

import logging
import re
import argparse
import sys
import json
import os
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Set, Any, Union
from dataclasses import dataclass, field
from enum import Enum, auto
from datetime import datetime
from difflib import SequenceMatcher
from io import BytesIO

# Third-party imports - with availability checks
try:
    from docx import Document
    from docx.shared import Pt
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from pyzbar.pyzbar import decode as decode_barcode
    PYZBAR_AVAILABLE = True
except ImportError:
    PYZBAR_AVAILABLE = False

# =============================================================================
# LOGGING SETUP
# =============================================================================

def setup_logging(level: int = logging.INFO) -> logging.Logger:
    """
    Configure logging for the artwork checker.
    
    Sets up console logging with timestamps and appropriate formatting
    for tracking the verification workflow.
    
    Args:
        level: Logging level (default: INFO)
        
    Returns:
        Configured logger instance
    """
    logging.basicConfig(
        level=level,
        format='%(asctime)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    return logging.getLogger(__name__)

logger = setup_logging()

# =============================================================================
# ENUMERATIONS
# =============================================================================

class ExtractionMethod(Enum):
    """
    Indicates how text was extracted from artwork.
    
    Used to track extraction confidence and determine
    when visual verification is required.
    """
    LIVE_TEXT = "live_text"          # Direct text extraction from PDF
    AI_PDF_COMPATIBLE = "ai_pdf"      # .ai file via PDF compatibility
    VISION_SYSTEM = "vision"          # AI vision fallback
    FAILED = "failed"                 # Extraction failed


class MatchType(Enum):
    """
    Classification of how well artwork text matches copy document.
    
    Used to determine status codes and visual verification requirements.
    """
    EXACT_MATCH = "exact"             # 100% character match
    NEAR_MATCH = "near"               # 95-99% match (minor differences)
    MISMATCH = "mismatch"             # <95% match (significant differences)
    MISSING_IN_ARTWORK = "missing"    # Copy exists, not found in artwork
    EXTRA_IN_ARTWORK = "extra"        # Found in artwork, not in copy
    REQUIRES_VERIFICATION = "verify"  # Needs visual confirmation


class StatusCode(Enum):
    """
    Status indicators for report findings.
    
    Maps to emoji indicators in the rendered output.
    """
    OK = "OK"       # ✅ Passed
    ATTN = "ATTN"   # ⚠️ Attention needed
    FAIL = "FAIL"   # ❌ Failed
    TBD = "TBD"     # Pending verification
    FYI = "FYI"     # Informational only


class RiskLevel(Enum):
    """
    Risk classification for regulatory claims.
    
    Determines escalation requirements for claims review.
    """
    LOW = "Low"
    MEDIUM = "Medium"
    HIGH = "High"


# =============================================================================
# CONFIGURATION
# =============================================================================

class Config:
    """
    Central configuration for the artwork checker.
    
    Contains all configurable parameters, thresholds, patterns,
    and rules used throughout the verification process.
    
    Attributes are organized by category for easy maintenance.
    
    Note:
        Modify these values to adjust checker behavior.
        Future Enhancement: Load from external config file.
    """
    
    # -------------------------------------------------------------------------
    # VERSION INFO
    # -------------------------------------------------------------------------
    VERSION = "2.0.0"
    BUILD_DATE = "2025-02"
    
    # -------------------------------------------------------------------------
    # MATCHING THRESHOLDS
    # -------------------------------------------------------------------------
    EXACT_MATCH_THRESHOLD = 100.0    # Must be 100% for exact match
    NEAR_MATCH_THRESHOLD = 95.0      # 95-99% considered near match
    MISMATCH_THRESHOLD = 95.0        # Below 95% is mismatch
    
    # -------------------------------------------------------------------------
    # ZOOM TRIGGERS
    # Conditions that trigger visual verification requirement
    # -------------------------------------------------------------------------
    ZOOM_FONT_SIZE_THRESHOLD = 6.5   # Font size <= this triggers zoom
    ZOOM_ON_NUMBERS = True           # Zoom if field contains numbers
    ZOOM_ON_PERCENTAGE = True        # Zoom if field contains %
    ZOOM_ON_DECIMALS = True          # Zoom if field contains decimals
    ZOOM_ON_UNITS = True             # Zoom if field contains units
    ZOOM_ON_NEGATION = True          # Zoom if field contains negation words
    ZOOM_CONFIDENCE_THRESHOLD = 100  # Zoom if confidence < 100%
    ZOOM_FUZZY_THRESHOLD = 100       # Zoom if fuzzy score < 100%
    
    # Unit patterns for zoom trigger
    UNIT_PATTERNS = [
        r'\b\d+\s*(mg|g|kg|ml|mL|ML|l|L|oz|OZ|fl\.?\s*oz|FL\.?\s*OZ|pt|qt|gal)\b',
        r'\b\d+\s*(mm|cm|m|in|inch|inches|ft|feet)\b',
        r'\bUS\s+FL\.?\s*OZ\.?\b',
    ]
    
    # Negation words for zoom trigger
    NEGATION_WORDS = [
        'no', 'not', 'free', 'only', 'without', 'never', 'none',
        'zero', 'moins', 'sans', 'non', 'aucun'  # French negations
    ]
    
    # -------------------------------------------------------------------------
    # CAPITALIZATION RULES
    # Fields that allow uppercase (all others must be lowercase)
    # -------------------------------------------------------------------------
    UPPERCASE_ALLOWED_FIELDS = [
        'Address Block',
        'Biorius Address', 
        'Formula Country of Origin',
        'Formula Country',
        'Ingredient List',
        'Fill Weight',
        'Net Weight',
    ]
    
    # Acronyms always allowed as uppercase anywhere
    ALLOWED_ACRONYMS = [
        'ML', 'mL', 'FL', 'OZ', 'USA', 'EU', 'UK', 'GB', 'BE', 'CA',
        'AHA', 'BHA', 'LHA', 'NP', 'SPF', 'UV', 'UVA', 'UVB',
        'PEG', 'PPG', 'EDTA', 'BHT', 'BHA',
        'NY', 'NYC', 'LA', 'SF',  # City abbreviations
        'LLC', 'INC', 'CO', 'LTD',  # Business abbreviations
    ]
    
    # INCI lowercase connector words (should NOT be capitalized)
    INCI_LOWERCASE_CONNECTORS = [
        'de', 'du', 'des', 'la', 'le', 'les', "d'", "l'",  # French
        'of', 'the', 'and', 'with',  # English
        'et', 'cum',  # Latin
    ]
    
    # -------------------------------------------------------------------------
    # EXCLUSION PATTERNS
    # Text patterns to exclude from analysis (non-product text)
    # -------------------------------------------------------------------------
    
    # Exact text matches to exclude (case-insensitive)
    EXCLUSION_EXACT_MATCHES = {
        # Color callouts
        'white', 'black', 'cyan', 'magenta', 'yellow',
        'c', 'm', 'y', 'k', 'cmyk', 'rgb',
        'spot gloss', 'matte', 'gloss', 'foil',
        
        # Packaging terms
        'bleed', 'trim', 'safe zone', 'die cut', 'dieline', 'die line',
        'fold', 'cut', 'glue', 'flap', 'score',
        'crop marks', 'fold here', 'cut here', 'glue flap',
        
        # Technical terms
        'proof', 'draft', 'final', 'approved',
        'fpo', 'for position only',
        'ol', 'outlined',
        
        # Orientation/view labels
        'front', 'back', 'side', 'top', 'bottom',
        'print side', 'inside', 'outside',
    }
    
    # Regex patterns to exclude
    EXCLUSION_REGEX_PATTERNS = [
        # Dimension patterns: 3+7/8, 16 3/32, 9.828
        r'^\d+\+?\d*/\d+$',
        r'^\d+\s+\d+/\d+$',
        r'^\d+\.\d{2,}$',
        
        # Reference numbers: IP250311C, VC_22640
        r'^IP\d+[A-Z]?$',
        r'^VC[_-]?\d+$',
        r'^ref\s*#?\s*:',
        r'^code\s*#?\s*:',
        
        # Pantone/PMS colors
        r'^PMS\s*\d+',
        r'^Pantone\s*\d+',
        r'^PMS\s+\w+\.\s*\w+',
        
        # Technical specifications
        r'^SCALE\s*:',
        r'^SIZE\s*:',
        r'^DATE\s*:',
        r'^DRAWN\s*:',
        r'^CHECKED\s*:',
        r'^APPROVED\s*:',
        r'^MATERIAL\s*:',
        r'^WEIGHT\s*:',
        r'^DWG\s*NO',
        r'^OFC\s*:',
        
        # Tolerances: X.X±0.5, .XX±0.012
        r'^[X\.]+±\d+',
        r'^\.\w+±\d+',
        r'^X°±\d+',
        r'GENERAL\s+TOLERANCES',
        r'MILLIMETRES',
        r'INCHES',
        
        # Company boilerplate
        r'INTEGRATED\s+PACKAGING',
        r'Kroger\s+Packaging',
        r'PROPRIETARY',
        r'CONFIDENTIAL',
        r'COPYRIGHT',
        r'This\s+design\s+concept\s+is\s+the\s+exclusive\s+property',
        r'all\s+rights\s+are\s+reserved',
        r'end\s+user\s*:',
        r'item\s*:',
        r'3RD\s+ANGLE\s+PROJECTION',
        
        # Label/packaging metadata
        r'^Label\s+Dieline',
        r'^Customer\s*:',
        r'^Diameter',
        r'^CLEARANCE\s+AREA',
        r'^Deco\s+Area',
        
        # Version/date stamps
        r'^\d{1,2}/\d{1,2}/\d{2,4}$',
        r'^\d{4}\.\d{2}\.\d{2}$',
    ]
    
    # Company names to exclude
    EXCLUSION_COMPANY_NAMES = [
        'INTEGRATED PACKAGING INDUSTRIES',
        'Kroger Packaging Inc',
        'Heat Makes Sense',
        'NVI',
    ]
    
    # -------------------------------------------------------------------------
    # INSTRUCTIONAL NOTE DETECTION
    # Patterns that indicate internal instructions vs actual copy
    # -------------------------------------------------------------------------
    INSTRUCTIONAL_PATTERNS = [
        # Yes/No starters
        r'^yes\s*[-–—]',
        r'^no\s*[-–—]',
        
        # Pending/placeholder indicators
        r'\bTBD\b',
        r'\bTBC\b',
        r'\bTODO\b',
        r'\bXXX+\b',
        r'\bPLACEHOLDER\b',
        r'\b[Pp]ending\b',
        
        # Production references
        r'\bPO\b',
        r'\bfirst\s+PO\b',
        r'\bproduction\b',
        
        # Conditional language
        r'if\s+it\s+can\s+fit',
        r'need\s+to\s+wait',
        r'waiting\s+for',
        r'once\s+confirmed',
        
        # Parenthetical instructions
        r'\(pending\)',
        r'\(see\s+attached\)',
        r'\(if\s+',
        r'\(TBD\)',
        r'\(optional\)',
        
        # Mentions and assignments
        r'@\w+',
        r'\b(?:ask|check\s+with|confirm\s+with)\b',
        
        # Question marks suggesting unresolved
        r'\?\s*$',
        
        # Date references in instruction context
        r'(?:updated|as\s+of|by)\s+\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
        
        # Regulatory/confirmation notes
        r'(?:Reg|regulatory)\s+confirmation',
        r'certification\s+(?:pending|needed)',
    ]
    
    # -------------------------------------------------------------------------
    # CLAIM RISK ASSESSMENT
    # Terms that indicate claims requiring substantiation
    # -------------------------------------------------------------------------
    HIGH_RISK_CLAIM_TERMS = [
        'clinically proven', 'clinically tested', 'clinically demonstrated',
        'dermatologist tested', 'dermatologist approved', 'dermatologist recommended',
        'doctor recommended', 'physician recommended',
        'scientifically proven', 'scientifically tested',
        'medically proven', 'FDA approved', 'FDA cleared',
        'patented', 'patent pending',
        'cliniquement prouvé', 'testé cliniquement',  # French
        'dermatologiquement testé',
    ]
    
    MEDIUM_RISK_CLAIM_TERMS = [
        'reduces', 'eliminates', 'removes', 'prevents',
        'anti-aging', 'anti-wrinkle', 'anti-acne',
        'repairs', 'restores', 'regenerates', 'renews',
        'strengthens', 'fortifies', 'rebuilds',
        'treats', 'heals', 'cures',
        'deeply moisturizes', 'intensely hydrates',
        'réduit', 'élimine', 'prévient',  # French
    ]
    
    # -------------------------------------------------------------------------
    # BARCODE CONFIGURATION
    # -------------------------------------------------------------------------
    BARCODE_RENDER_DPI = 4  # Multiplier for PDF rendering
    BARCODE_TYPES = {
        'UPC-A': 12,
        'UPC-E': 8,
        'EAN-13': 13,
        'EAN-8': 8,
    }
    
    # -------------------------------------------------------------------------
    # SNAPSHOT CONFIGURATION
    # -------------------------------------------------------------------------
    SNAPSHOT_ENABLED = True
    SNAPSHOT_MODE = 'full_page'  # 'full_page' or 'crop' (crop is placeholder)
    SNAPSHOT_DPI = 150           # DPI for snapshot rendering
    SNAPSHOT_MAX_PER_PAGE = 20   # Max annotations per page
    
    # Annotation colors (RGB)
    ANNOTATION_COLORS = {
        'FAIL': (255, 0, 0),      # Red
        'ATTN': (255, 165, 0),    # Orange
        'OK': (0, 255, 0),        # Green (not typically annotated)
        'TBD': (128, 128, 128),   # Gray
    }
    ANNOTATION_LINE_WIDTH = 3
    ANNOTATION_FONT_SIZE = 14
    
    # -------------------------------------------------------------------------
    # OUTPUT CONFIGURATION
    # -------------------------------------------------------------------------
    STATUS_EMOJI = {
        'OK': '✅',
        'ATTN': '⚠️',
        'FAIL': '❌',
        'TBD': '🔍',
        'FYI': 'ℹ️',
    }
    
    # -------------------------------------------------------------------------
    # FIELD PARSING CONFIGURATION
    # -------------------------------------------------------------------------
    # Fields that should be split into sub-fields by paragraph
    PARAGRAPH_SPLIT_FIELDS = [
        'Marketing + Usage Copy',
        'Marketing Copy',
        'Usage Copy',
        'Hero Ingredient Call-outs',
        'Hero Ingredients',
        'Key Ingredients',
        'Pack Claims',
    ]
    
    # Sub-field naming templates
    SUBFIELD_TEMPLATES = {
        'Marketing + Usage Copy': ['Marketing Copy', 'Scent Line', 'Usage Instructions'],
        'Hero Ingredient Call-outs': ['Hero Ingredient {}'],
        'Pack Claims': ['Pack Claim {}'],
    }


# Global config instance
config = Config()


# =============================================================================
# DATA MODELS
# =============================================================================

@dataclass
class TextRun:
    """
    A single extracted text element from artwork with metadata.
    
    Represents one contiguous piece of text found during extraction,
    along with its location, font information, and extraction confidence.
    
    Attributes:
        text: The extracted text content
        page_number: 1-based page number where text was found
        bbox: Bounding box coordinates (x0, y0, x1, y1)
        font_name: Name of the font used
        font_size: Font size in points
        extraction_method: How the text was extracted
        confidence: Confidence score (0.0 to 1.0)
        rotation: Text rotation in degrees (0, 90, 180, 270)
    """
    text: str
    page_number: int
    bbox: Optional[Tuple[float, float, float, float]] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    extraction_method: ExtractionMethod = ExtractionMethod.LIVE_TEXT
    confidence: float = 1.0
    rotation: int = 0


@dataclass
class CopyField:
    """
    A field from the copy document to be verified.
    
    Represents a single piece of copy text that should appear
    on the artwork, with its location and language context.
    
    Attributes:
        field_name: Name of the field (e.g., "Product Name")
        panel: Location on artwork (e.g., "Front Panel", "Back Panel")
        language: Language code (e.g., "EN", "FR")
        text: The text content from copy document
        is_strikethrough: Whether text had strikethrough formatting
        is_legacy: Whether text is legacy (before ` -> ` separator)
        parent_field: Original field name if this is a sub-field
        subfield_index: Index if this is part of a split field
    """
    field_name: str
    panel: str
    language: str
    text: str
    is_strikethrough: bool = False
    is_legacy: bool = False
    parent_field: Optional[str] = None
    subfield_index: Optional[int] = None


@dataclass
class MatchFinding:
    """
    Result of comparing a copy field against artwork.
    
    Contains all information about whether the copy text was found
    in the artwork and how well it matched.
    
    Attributes:
        field_name: Name of the field being checked
        panel: Panel location on artwork
        language: Language code
        copy_value: Text from copy document
        artwork_value: Text found in artwork (or None)
        match_type: Classification of match quality
        similarity_score: Fuzzy match score (0-100)
        status_code: Overall status (OK, ATTN, FAIL, etc.)
        notes: List of notes/warnings for this finding
        requires_zoom: Whether visual verification is required
        zoom_reasons: Why zoom was triggered
        issue_id: Unique identifier for this issue (e.g., "D-001")
        bbox: Bounding box of matched text in artwork
    """
    field_name: str
    panel: str
    language: str
    copy_value: str
    artwork_value: Optional[str]
    match_type: MatchType
    similarity_score: float
    status_code: StatusCode
    notes: List[str] = field(default_factory=list)
    requires_zoom: bool = False
    zoom_reasons: List[str] = field(default_factory=list)
    issue_id: Optional[str] = None
    bbox: Optional[Tuple[float, float, float, float]] = None


@dataclass
class CopyQualityIssue:
    """
    An issue found in the copy document itself.
    
    Section 3A findings - issues with the copy document text
    independent of artwork comparison.
    
    Attributes:
        language: Language code
        field_name: Name of the field with issue
        original_text: The problematic text
        issue_type: Category of issue (e.g., "Capitalization", "Typo")
        recommendation: Suggested fix
        status_code: Severity (OK, ATTN, FAIL)
    """
    language: str
    field_name: str
    original_text: str
    issue_type: str
    recommendation: str
    status_code: StatusCode


@dataclass 
class ClaimRisk:
    """
    Risk assessment for a regulatory claim.
    
    Section 3B findings - claims that may require substantiation
    or regulatory review.
    
    Attributes:
        language: Language code
        claim_text: The claim being assessed
        risk_level: LOW, MEDIUM, or HIGH
        rationale: Why this risk level was assigned
        regions: Regions where this claim may be problematic
        recommended_action: What to do (Keep, Modify, Escalate)
        status_code: Status indicator
    """
    language: str
    claim_text: str
    risk_level: RiskLevel
    rationale: str
    regions: List[str]
    recommended_action: str
    status_code: StatusCode


@dataclass
class ConversionCheck:
    """
    Volume/weight conversion verification.
    
    Section 3C findings - verifying mL to fl oz conversions
    are mathematically correct within tolerance.
    
    Attributes:
        source_field: Field where declaration was found
        declared_ml: Declared metric value
        declared_floz: Declared imperial value  
        calculated_floz: Calculated imperial value
        within_tolerance: Whether within ±0.10 fl oz
        status_code: Pass/fail status
        notes: Additional notes
    """
    source_field: str
    declared_ml: float
    declared_floz: float
    calculated_floz: float
    within_tolerance: bool
    status_code: StatusCode
    notes: str = ""


@dataclass
class FontMeasurement:
    """
    Font size measurement from artwork.
    
    Section 3E findings - tracking font sizes for
    regulatory compliance checking.
    
    Attributes:
        text: Sample text with this font size
        page_number: Page where found
        font_name: Font name
        font_size_pt: Size in points
        bbox: Location in artwork
    """
    text: str
    page_number: int
    font_name: str
    font_size_pt: float
    bbox: Optional[Tuple[float, float, float, float]] = None


@dataclass
class BarcodeResult:
    """
    Result of barcode scanning and validation.
    
    Section 3F findings - barcode quality and digit verification.
    
    Attributes:
        symbology: Type of barcode (UPC-A, EAN-13, etc.)
        decoded_digits: Digits read from barcode scan
        printed_digits: Digits printed below barcode
        digits_match: Whether decoded matches printed
        check_digit_valid: Whether check digit is correct
        scan_successful: Whether barcode was scannable
        quality_notes: Notes about scan quality
        status_code: Overall status
    """
    symbology: str
    decoded_digits: Optional[str]
    printed_digits: Optional[str]
    digits_match: bool
    check_digit_valid: bool
    scan_successful: bool
    quality_notes: str
    status_code: StatusCode


@dataclass
class SnapshotAnnotation:
    """
    Annotation to be drawn on a snapshot image.
    
    Represents a single issue to highlight on the
    full-page annotated snapshot.
    
    Attributes:
        issue_id: Reference ID (e.g., "D-001")
        bbox: Bounding box to highlight
        status: Status code for color selection
        page_number: Page this annotation belongs to
    """
    issue_id: str
    bbox: Tuple[float, float, float, float]
    status: StatusCode
    page_number: int


@dataclass
class ExtractionResult:
    """
    Complete extraction result from an artwork file.
    
    Contains all text runs extracted from the artwork
    along with metadata about the extraction process.
    
    Attributes:
        file_path: Path to the artwork file
        file_type: Type of file (pdf, ai)
        extraction_method: Primary method used
        confidence: Overall extraction confidence
        text_runs: All extracted text elements
        pages_processed: Number of pages processed
        pages_with_text: Pages that contained text
        warnings: Any warnings during extraction
        metadata: Additional file metadata
    """
    file_path: str
    file_type: str
    extraction_method: ExtractionMethod
    confidence: float
    text_runs: List[TextRun] = field(default_factory=list)
    pages_processed: int = 0
    pages_with_text: int = 0
    warnings: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class CopyDocument:
    """
    Parsed copy document data structure.
    
    Contains all fields extracted from the Word copy document,
    organized by panel and language.
    
    Attributes:
        file_path: Path to the copy document
        panels: Nested dict: panel -> language -> field_name -> text
        fields: Flat list of all CopyField objects
        instructional_notes: Fields flagged as internal instructions
        strikethrough_text: Text that had strikethrough formatting
        legacy_text: Text before ` -> ` separators
        warnings: Extraction warnings
        metadata: File metadata
    """
    file_path: str
    panels: Dict[str, Dict[str, Dict[str, str]]] = field(default_factory=dict)
    fields: List[CopyField] = field(default_factory=list)
    instructional_notes: List[Tuple[str, str, str]] = field(default_factory=list)
    strikethrough_text: List[Tuple[str, str, str]] = field(default_factory=list)
    legacy_text: List[Tuple[str, str, str]] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def fuzzy_ratio(str1: str, str2: str) -> float:
    """
    Calculate fuzzy similarity ratio between two strings.
    
    Uses Python's built-in SequenceMatcher for reliable comparison
    without external dependencies.
    
    Args:
        str1: First string to compare
        str2: Second string to compare
        
    Returns:
        Similarity score from 0.0 to 100.0
        
    Example:
        >>> fuzzy_ratio("hello world", "hello world")
        100.0
        >>> fuzzy_ratio("hello", "hallo")
        80.0
    """
    if not str1 and not str2:
        return 100.0
    if not str1 or not str2:
        return 0.0
    return SequenceMatcher(None, str1, str2).ratio() * 100


def sanitize_for_markdown(text: str) -> str:
    """
    Sanitize text for safe inclusion in markdown tables.
    
    Removes or escapes characters that would break markdown
    table formatting.
    
    Args:
        text: Raw text to sanitize
        
    Returns:
        Sanitized text safe for markdown tables
    """
    if not text:
        return ""
    
    text = str(text)
    
    # Replace newlines with space (critical for table cells)
    text = text.replace('\n', ' ')
    text = text.replace('\r', '')
    text = text.replace('\t', ' ')
    
    # Escape pipe characters
    text = text.replace('|', '\\|')
    
    # Collapse multiple spaces
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()


def truncate_text(text: str, max_length: int = 100, suffix: str = "...") -> str:
    """
    Truncate text to maximum length with suffix.
    
    Preserves word boundaries when possible.
    
    Args:
        text: Text to truncate
        max_length: Maximum length including suffix
        suffix: String to append when truncated
        
    Returns:
        Truncated text with suffix if needed
    """
    if not text or len(text) <= max_length:
        return text or ""
    
    # Try to break at word boundary
    truncated = text[:max_length - len(suffix)]
    last_space = truncated.rfind(' ')
    
    if last_space > max_length * 0.5:  # At least half the length
        truncated = truncated[:last_space]
    
    return truncated + suffix


# =============================================================================
# TEXT NORMALIZATION
# =============================================================================

class TextNormalizer:
    """
    Normalize text for comparison while preserving meaningful differences.
    
    Handles character-by-character comparison requirements while
    normalizing acceptable variations (whitespace, quote styles, etc.).
    
    The goal is to compare semantic content while allowing for
    minor typographic variations that don't affect meaning.
    """
    
    @staticmethod
    def normalize(text: str) -> str:
        """
        Normalize text for comparison.
        
        Applies the following normalizations:
        - Collapse multiple whitespace to single space
        - Normalize various quote characters to standard quotes
        - Normalize various dash characters to standard hyphen/em-dash
        - Normalize apostrophes
        - Normalize non-breaking spaces
        - Strip leading/trailing whitespace
        
        Args:
            text: Raw text to normalize
            
        Returns:
            Normalized text for comparison
            
        Note:
            Does NOT change case - case sensitivity is preserved
            for character-by-character comparison.
        """
        if not text:
            return ""
        
        normalized = text
        
        # Normalize whitespace (but preserve case)
        normalized = re.sub(r'\s+', ' ', normalized)
        normalized = normalized.strip()
        
        # Normalize quotes
        # Straight quotes (curly to straight)
        normalized = normalized.replace('\u201C', '"').replace('\u201D', '"')  # double quotes
        normalized = normalized.replace('\u2018', "'").replace('\u2019', "'")  # single quotes
        # Guillemets to straight quotes (for comparison only)
        normalized = normalized.replace('\u00AB', '"').replace('\u00BB', '"')
        
        # Normalize dashes
        normalized = normalized.replace('\u2013', '-')  # en-dash to hyphen
        normalized = normalized.replace('\u2014', '\u2014')  # preserve em-dash
        normalized = normalized.replace('\u2010', '-')  # unicode hyphen
        
        # Normalize apostrophes
        normalized = normalized.replace('\u2019', "'")
        normalized = normalized.replace('\u02BC', "'")
        
        # Normalize non-breaking spaces
        normalized = normalized.replace('\u00A0', ' ')
        normalized = normalized.replace('\u202F', ' ')
        
        # Normalize ligatures
        normalized = normalized.replace('\uFB01', 'fi')
        normalized = normalized.replace('\uFB02', 'fl')
        normalized = normalized.replace('\uFB00', 'ff')
        normalized = normalized.replace('\uFB03', 'ffi')
        normalized = normalized.replace('\uFB04', 'ffl')
        
        return normalized
    
    @staticmethod
    def normalize_for_search(text: str) -> str:
        """
        Aggressive normalization for text searching/matching.
        
        Used when trying to find text in artwork - more lenient
        than exact comparison normalization.
        
        Args:
            text: Text to normalize for searching
            
        Returns:
            Lowercased, heavily normalized text
        """
        if not text:
            return ""
        
        normalized = TextNormalizer.normalize(text)
        normalized = normalized.lower()
        
        # Remove all punctuation for search
        normalized = re.sub(r'[^\w\s]', '', normalized)
        normalized = re.sub(r'\s+', ' ', normalized)
        
        return normalized.strip()


# =============================================================================
# ZOOM TRIGGER DETECTION
# =============================================================================

class ZoomTriggerDetector:
    """
    Determines when visual verification (zoom) is required.
    
    Based on content characteristics that increase risk of
    extraction errors or require higher precision verification.
    
    Uses the zoom trigger rules defined in Config.
    """
    
    @staticmethod
    def check_triggers(
        text: str,
        font_size: Optional[float] = None,
        confidence: float = 100.0,
        fuzzy_score: float = 100.0
    ) -> Tuple[bool, List[str]]:
        """
        Check if any zoom triggers are activated for given text.
        
        Args:
            text: The text content to check
            font_size: Font size in points (if known)
            confidence: Extraction confidence (0-100)
            fuzzy_score: Fuzzy match score (0-100)
            
        Returns:
            Tuple of (requires_zoom: bool, reasons: List[str])
            
        Example:
            >>> requires, reasons = ZoomTriggerDetector.check_triggers("250 mL", 5.5)
            >>> print(requires)
            True
            >>> print(reasons)
            ['Font size 5.5pt ≤ 6.5pt threshold', 'Contains units']
        """
        triggers = []
        
        # Check font size
        if font_size is not None and font_size <= config.ZOOM_FONT_SIZE_THRESHOLD:
            triggers.append(f"Font size {font_size}pt ≤ {config.ZOOM_FONT_SIZE_THRESHOLD}pt threshold")
        
        # Check confidence
        if confidence < config.ZOOM_CONFIDENCE_THRESHOLD:
            triggers.append(f"Extraction confidence {confidence:.0f}% < {config.ZOOM_CONFIDENCE_THRESHOLD}%")
        
        # Check fuzzy score
        if fuzzy_score < config.ZOOM_FUZZY_THRESHOLD:
            triggers.append(f"Fuzzy match {fuzzy_score:.1f}% < {config.ZOOM_FUZZY_THRESHOLD}%")
        
        if text:
            # Check for numbers
            if config.ZOOM_ON_NUMBERS and re.search(r'\d', text):
                triggers.append("Contains numbers")
            
            # Check for percentage
            if config.ZOOM_ON_PERCENTAGE and '%' in text:
                triggers.append("Contains percentage")
            
            # Check for decimals
            if config.ZOOM_ON_DECIMALS and re.search(r'\d+\.\d+', text):
                triggers.append("Contains decimal numbers")
            
            # Check for units
            if config.ZOOM_ON_UNITS:
                for pattern in config.UNIT_PATTERNS:
                    if re.search(pattern, text, re.IGNORECASE):
                        triggers.append("Contains units")
                        break
            
            # Check for negation words
            if config.ZOOM_ON_NEGATION:
                text_lower = text.lower()
                for negation in config.NEGATION_WORDS:
                    if re.search(rf'\b{re.escape(negation)}\b', text_lower):
                        triggers.append(f"Contains negation word: '{negation}'")
                        break
        
        return len(triggers) > 0, triggers


# =============================================================================
# EXCLUSION PATTERN CHECKER
# =============================================================================

class ExclusionChecker:
    """
    Checks if text should be excluded from analysis.
    
    Identifies non-product text like technical notes, dimensions,
    color callouts, and packaging metadata that should not be
    compared against the copy document.
    """
    
    @staticmethod
    def should_exclude(text: str) -> Tuple[bool, Optional[str]]:
        """
        Check if text should be excluded from analysis.
        
        Args:
            text: Text to check
            
        Returns:
            Tuple of (should_exclude: bool, reason: Optional[str])
            
        Example:
            >>> exclude, reason = ExclusionChecker.should_exclude("PMS 225c")
            >>> print(exclude)
            True
            >>> print(reason)
            "Matches exclusion pattern: Pantone color"
        """
        if not text or len(text.strip()) < 2:
            return True, "Empty or too short"
        
        text_clean = text.strip()
        text_lower = text_clean.lower()
        
        # Check exact matches
        if text_lower in config.EXCLUSION_EXACT_MATCHES:
            return True, f"Exact match exclusion: '{text_clean}'"
        
        # Check company names
        for company in config.EXCLUSION_COMPANY_NAMES:
            if company.lower() in text_lower:
                return True, f"Company name: '{company}'"
        
        # Check regex patterns
        for pattern in config.EXCLUSION_REGEX_PATTERNS:
            if re.search(pattern, text_clean, re.IGNORECASE):
                return True, f"Pattern match: '{pattern[:30]}...'"
        
        return False, None
    
    @staticmethod
    def filter_text_runs(text_runs: List[TextRun]) -> List[TextRun]:
        """
        Filter out text runs that should be excluded.
        
        Args:
            text_runs: List of extracted text runs
            
        Returns:
            Filtered list with exclusions removed
        """
        filtered = []
        excluded_count = 0
        
        for run in text_runs:
            should_exclude, reason = ExclusionChecker.should_exclude(run.text)
            if not should_exclude:
                filtered.append(run)
            else:
                excluded_count += 1
                logger.debug(f"Excluded text: '{run.text[:50]}' - {reason}")
        
        if excluded_count > 0:
            logger.info(f"Excluded {excluded_count} non-product text elements")
        
        return filtered


# =============================================================================
# INSTRUCTIONAL NOTE DETECTOR
# =============================================================================

class InstructionalNoteDetector:
    """
    Detects internal instructions vs actual copy text.
    
    Identifies text that appears to be internal notes, placeholders,
    or instructions that should not appear on final artwork.
    """
    
    @staticmethod
    def is_instructional(text: str) -> Tuple[bool, Optional[str]]:
        """
        Check if text appears to be an internal instruction.
        
        Args:
            text: Text to check
            
        Returns:
            Tuple of (is_instruction: bool, pattern_matched: Optional[str])
            
        Example:
            >>> is_instr, pattern = InstructionalNoteDetector.is_instructional(
            ...     "yes – also add B Corp logo"
            ... )
            >>> print(is_instr)
            True
        """
        if not text:
            return False, None
        
        for pattern in config.INSTRUCTIONAL_PATTERNS:
            if re.search(pattern, text, re.IGNORECASE):
                return True, pattern
        
        # Check for empty placeholder fields
        if text.strip() in ['', '-', 'N/A', 'n/a', 'NA', 'na', '—', '–']:
            return True, "Empty/placeholder value"
        
        return False, None



# =============================================================================
# INCI CAPITALIZATION VALIDATOR
# =============================================================================

class INCIValidator:
    """
    Validates INCI (International Nomenclature of Cosmetic Ingredients) capitalization.
    
    INCI names follow specific capitalization rules:
    - First letter of each word capitalized
    - Except connector words (de, of, the, etc.)
    - Parenthetical common names follow similar rules
    
    Future Enhancement:
        Add connection to external INCI database API for
        comprehensive validation of ingredient names.
    """
    
    @staticmethod
    def validate_inci_capitalization(text: str) -> Tuple[bool, List[str]]:
        """
        Validate INCI capitalization in ingredient text.
        
        Args:
            text: Ingredient list text to validate
            
        Returns:
            Tuple of (is_valid: bool, issues: List[str])
            
        Example:
            >>> valid, issues = INCIValidator.validate_inci_capitalization(
            ...     "cocos nucifera (Coconut) Oil"
            ... )
            >>> print(valid)
            False
            >>> print(issues)
            ["'cocos' should be 'Cocos'", "'nucifera' should be 'Nucifera'"]
        """
        if not text:
            return True, []
        
        issues = []
        
        # Split into words, preserving parentheses
        words = re.findall(r'\(?\b\w+\b\)?', text)
        
        for word in words:
            # Remove parentheses for checking
            clean_word = word.strip('()')
            
            # Skip if it's a connector word
            if clean_word.lower() in config.INCI_LOWERCASE_CONNECTORS:
                # Connector should be lowercase
                if clean_word[0].isupper() and clean_word.lower() in ['de', 'du', 'des', 'la', 'le', 'les', 'of', 'the', 'and', 'with', 'et', 'cum']:
                    issues.append(f"'{clean_word}' should be lowercase '{clean_word.lower()}'")
                continue
            
            # Skip numbers and very short words
            if len(clean_word) < 2 or clean_word.isdigit():
                continue
            
            # Skip if all uppercase (likely acronym)
            if clean_word.isupper():
                continue
            
            # Skip if it's a known acronym
            if clean_word.upper() in config.ALLOWED_ACRONYMS:
                continue
            
            # Check first letter capitalization
            if clean_word[0].islower() and clean_word[0].isalpha():
                correct = clean_word[0].upper() + clean_word[1:]
                issues.append(f"'{clean_word}' should be '{correct}'")
        
        return len(issues) == 0, issues
    
    @staticmethod
    def validate_ingredient_list(ingredient_text: str) -> List[Dict[str, str]]:
        """
        Validate a full ingredient list and return detailed findings.
        
        Args:
            ingredient_text: Full ingredient list text
            
        Returns:
            List of issue dictionaries with details
            
        Note:
            Future Enhancement: Connect to INCI database API
            for comprehensive ingredient validation.
        """
        findings = []
        
        # Split by common ingredient separators
        ingredients = re.split(r'[,;]', ingredient_text)
        
        for ingredient in ingredients:
            ingredient = ingredient.strip()
            if not ingredient:
                continue
            
            is_valid, issues = INCIValidator.validate_inci_capitalization(ingredient)
            
            if not is_valid:
                findings.append({
                    'ingredient': ingredient[:50],
                    'issues': issues,
                    'severity': 'ATTN'
                })
        
        return findings


# =============================================================================
# COPY QUALITY CHECKER
# =============================================================================

class CopyQualityChecker:
    """
    Checks copy document text for quality issues.
    
    Section 3A - Analyzes copy text for:
    - Typography and punctuation errors
    - Capitalization rule violations
    - Spelling issues
    - Formatting problems
    
    Note:
        This only checks the copy document itself,
        not the comparison with artwork.
    """
    
    @staticmethod
    def check_capitalization(
        text: str, 
        field_name: str,
        language: str
    ) -> List[CopyQualityIssue]:
        """
        Check text for capitalization rule violations.
        
        All text should be lowercase except for specific fields
        and acronyms.
        
        Args:
            text: Text to check
            field_name: Name of the field
            language: Language code
            
        Returns:
            List of capitalization issues found
        """
        issues = []
        
        if not text:
            return issues
        
        # Check if this field allows uppercase
        allows_uppercase = any(
            allowed.lower() in field_name.lower() 
            for allowed in config.UPPERCASE_ALLOWED_FIELDS
        )
        
        if allows_uppercase:
            return issues  # Skip capitalization check for allowed fields
        
        # Find words that are incorrectly capitalized
        words = re.findall(r'\b[A-Za-z]+\b', text)
        
        for word in words:
            # Skip if it's an allowed acronym
            if word.upper() in config.ALLOWED_ACRONYMS:
                continue
            
            # Skip single letters
            if len(word) <= 1:
                continue
            
            # Check if word has unexpected capitalization
            # First letter uppercase when it shouldn't be (not sentence start)
            if word[0].isupper() and not word.isupper():
                # Check if it's at the start of text or after period
                word_pos = text.find(word)
                if word_pos > 0:
                    # Check what's before this word
                    before = text[:word_pos].rstrip()
                    if before and before[-1] not in '.!?':
                        issues.append(CopyQualityIssue(
                            language=language,
                            field_name=field_name,
                            original_text=word,
                            issue_type="Capitalization",
                            recommendation=f"'{word}' should be lowercase '{word.lower()}'",
                            status_code=StatusCode.FAIL
                        ))
        
        return issues
    
    @staticmethod
    def check_punctuation(
        text: str,
        field_name: str,
        language: str
    ) -> List[CopyQualityIssue]:
        """
        Check text for punctuation issues.
        
        Args:
            text: Text to check
            field_name: Name of the field
            language: Language code
            
        Returns:
            List of punctuation issues found
        """
        issues = []
        
        if not text:
            return issues
        
        # Double punctuation
        if re.search(r'[.]{2,}(?!\.)', text):  # Allow ellipsis
            issues.append(CopyQualityIssue(
                language=language,
                field_name=field_name,
                original_text=text[:50],
                issue_type="Punctuation",
                recommendation="Check for double periods",
                status_code=StatusCode.ATTN
            ))
        
        if ',,' in text:
            issues.append(CopyQualityIssue(
                language=language,
                field_name=field_name,
                original_text=text[:50],
                issue_type="Punctuation",
                recommendation="Remove double commas",
                status_code=StatusCode.ATTN
            ))
        
        # Multiple spaces
        if '  ' in text:
            issues.append(CopyQualityIssue(
                language=language,
                field_name=field_name,
                original_text=text[:50],
                issue_type="Formatting",
                recommendation="Remove extra spaces",
                status_code=StatusCode.ATTN
            ))
        
        # French-specific: space before colon/semicolon
        if language == 'FR':
            if re.search(r'\S[:;]', text):  # No space before : or ;
                # This is actually incorrect for French
                pass  # French typically needs space before these
        
        return issues
    
    @staticmethod
    def check_legacy_arrow(
        text: str,
        field_name: str,
        language: str
    ) -> List[CopyQualityIssue]:
        """
        Check for legacy text markers (` -> ` separator).
        
        Args:
            text: Text to check
            field_name: Name of the field
            language: Language code
            
        Returns:
            List of legacy text issues found
        """
        issues = []
        
        if ' -> ' in text:
            issues.append(CopyQualityIssue(
                language=language,
                field_name=field_name,
                original_text=text[:100],
                issue_type="Legacy Content",
                recommendation="Contains ` -> ` separator - old text should be removed",
                status_code=StatusCode.ATTN
            ))
        
        return issues
    
    @staticmethod
    def check_instructional_notes(
        text: str,
        field_name: str,
        language: str
    ) -> List[CopyQualityIssue]:
        """
        Check for internal instructional notes that shouldn't be in final copy.
        
        Args:
            text: Text to check
            field_name: Name of the field
            language: Language code
            
        Returns:
            List of instructional note issues found
        """
        issues = []
        
        is_instruction, pattern = InstructionalNoteDetector.is_instructional(text)
        
        if is_instruction:
            issues.append(CopyQualityIssue(
                language=language,
                field_name=field_name,
                original_text=text[:100],
                issue_type="Instructional Note",
                recommendation=f"Remove internal instruction (matched: {pattern[:30] if pattern else 'N/A'})",
                status_code=StatusCode.FAIL
            ))
        
        return issues
    
    @classmethod
    def analyze_copy_document(
        cls,
        copy_doc: CopyDocument
    ) -> List[CopyQualityIssue]:
        """
        Perform complete quality analysis on copy document.
        
        Args:
            copy_doc: Parsed copy document
            
        Returns:
            List of all quality issues found
        """
        all_issues = []
        
        for field in copy_doc.fields:
            # Skip legacy/strikethrough text
            if field.is_legacy or field.is_strikethrough:
                continue
            
            # Run all checks
            all_issues.extend(cls.check_capitalization(
                field.text, field.field_name, field.language
            ))
            all_issues.extend(cls.check_punctuation(
                field.text, field.field_name, field.language
            ))
            all_issues.extend(cls.check_legacy_arrow(
                field.text, field.field_name, field.language
            ))
            all_issues.extend(cls.check_instructional_notes(
                field.text, field.field_name, field.language
            ))
        
        # Check INCI in ingredient fields
        for field in copy_doc.fields:
            if 'ingredient' in field.field_name.lower():
                inci_findings = INCIValidator.validate_ingredient_list(field.text)
                for finding in inci_findings:
                    all_issues.append(CopyQualityIssue(
                        language=field.language,
                        field_name=field.field_name,
                        original_text=finding['ingredient'],
                        issue_type="INCI Capitalization",
                        recommendation="; ".join(finding['issues']),
                        status_code=StatusCode.ATTN
                    ))
        
        logger.info(f"Copy quality analysis found {len(all_issues)} issues")
        return all_issues


# =============================================================================
# CLAIM RISK ASSESSOR
# =============================================================================

class ClaimRiskAssessor:
    """
    Assesses regulatory risk of product claims.
    
    Section 3B - Evaluates claims for:
    - High-risk terms requiring substantiation
    - Medium-risk therapeutic/efficacy claims
    - Regional regulatory concerns
    """
    
    @staticmethod
    def assess_claim(claim_text: str, language: str) -> ClaimRisk:
        """
        Assess risk level of a single claim.
        
        Args:
            claim_text: The claim text to assess
            language: Language code
            
        Returns:
            ClaimRisk object with assessment details
        """
        claim_lower = claim_text.lower()
        
        # Check high-risk terms
        for term in config.HIGH_RISK_CLAIM_TERMS:
            if term.lower() in claim_lower:
                return ClaimRisk(
                    language=language,
                    claim_text=claim_text,
                    risk_level=RiskLevel.HIGH,
                    rationale=f"Contains high-risk term: '{term}' - requires substantiation",
                    regions=["USA", "EU", "UK", "CA"],
                    recommended_action="Escalate",
                    status_code=StatusCode.ATTN
                )
        
        # Check medium-risk terms
        for term in config.MEDIUM_RISK_CLAIM_TERMS:
            if term.lower() in claim_lower:
                return ClaimRisk(
                    language=language,
                    claim_text=claim_text,
                    risk_level=RiskLevel.MEDIUM,
                    rationale=f"Contains efficacy claim: '{term}' - verify support",
                    regions=["USA", "EU", "UK"],
                    recommended_action="Verify",
                    status_code=StatusCode.ATTN
                )
        
        # Default to low risk
        return ClaimRisk(
            language=language,
            claim_text=claim_text,
            risk_level=RiskLevel.LOW,
            rationale="Cosmetic/descriptive claim - acceptable",
            regions=["All"],
            recommended_action="Keep",
            status_code=StatusCode.OK
        )
    
    @classmethod
    def assess_all_claims(cls, copy_doc: CopyDocument) -> List[ClaimRisk]:
        """
        Assess all claims in a copy document.
        
        Args:
            copy_doc: Parsed copy document
            
        Returns:
            List of ClaimRisk assessments
        """
        assessments = []
        
        for field in copy_doc.fields:
            # Only check claim fields
            if 'claim' not in field.field_name.lower():
                continue
            
            # Skip legacy/strikethrough
            if field.is_legacy or field.is_strikethrough:
                continue
            
            # Assess each claim (handle multi-line)
            claims = field.text.split('\n')
            for claim in claims:
                claim = claim.strip()
                if claim:
                    assessment = cls.assess_claim(claim, field.language)
                    assessments.append(assessment)
        
        logger.info(f"Assessed {len(assessments)} claims")
        return assessments




# =============================================================================
# COPY DOCUMENT EXTRACTOR
# =============================================================================

class CopyDocumentExtractor:
    """
    Extracts and parses copy documents (.docx files).
    
    Handles:
    - Table-based copy document structure
    - Strikethrough text detection
    - Legacy text (` -> ` separator) parsing
    - Multi-language extraction
    - Field splitting for multi-paragraph content
    
    The copy document is expected to have a specific structure
    with panels (Front/Back) and language columns.
    """
    
    def __init__(self):
        """Initialize the extractor."""
        self.warnings = []
    
    def extract(self, file_path: Path) -> CopyDocument:
        """
        Extract all data from a copy document.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            CopyDocument with all extracted data
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is invalid
        """
        self.warnings = []
        
        if not PYTHON_DOCX_AVAILABLE:
            self.warnings.append("python-docx not available")
            return CopyDocument(file_path=str(file_path), warnings=self.warnings)
        
        try:
            doc = Document(file_path)
            logger.info(f"Opened copy document: {file_path.name}")
            
            # Extract panels data from tables
            panels_data = self._extract_from_tables(doc)
            
            # Build flat field list
            fields = self._build_field_list(panels_data, doc)
            
            # Detect instructional notes
            instructional_notes = []
            for field in fields:
                is_instr, pattern = InstructionalNoteDetector.is_instructional(field.text)
                if is_instr:
                    instructional_notes.append((field.field_name, field.language, field.text[:100]))
            
            # Get strikethrough and legacy text
            strikethrough = [(f.field_name, f.language, f.text[:100]) 
                           for f in fields if f.is_strikethrough]
            legacy = [(f.field_name, f.language, f.text[:100]) 
                     for f in fields if f.is_legacy]
            
            # Remove legacy/strikethrough from active fields
            active_fields = [f for f in fields if not f.is_legacy and not f.is_strikethrough]
            
            return CopyDocument(
                file_path=str(file_path),
                panels=panels_data,
                fields=active_fields,
                instructional_notes=instructional_notes,
                strikethrough_text=strikethrough,
                legacy_text=legacy,
                warnings=self.warnings,
                metadata={
                    'tables_count': len(doc.tables),
                    'total_fields': len(active_fields),
                    'languages_found': list(set(f.language for f in active_fields))
                }
            )
            
        except Exception as e:
            self.warnings.append(f"Extraction error: {str(e)}")
            logger.error(f"Failed to extract copy document: {e}")
            return CopyDocument(file_path=str(file_path), warnings=self.warnings)
    
    def _extract_from_tables(self, doc: Document) -> Dict[str, Dict[str, Dict[str, str]]]:
        """
        Extract structured data from document tables.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Nested dict: panel -> language -> field_name -> text
        """
        panels_data = {}
        
        for table_idx, table in enumerate(doc.tables):
            try:
                table_data = self._parse_copy_table(table, doc)
                
                # Merge into panels_data
                for panel, languages in table_data.items():
                    if panel not in panels_data:
                        panels_data[panel] = {}
                    for lang, fields in languages.items():
                        if lang not in panels_data[panel]:
                            panels_data[panel][lang] = {}
                        panels_data[panel][lang].update(fields)
                        
            except Exception as e:
                self.warnings.append(f"Error parsing table {table_idx + 1}: {str(e)}")
                logger.warning(f"Table {table_idx + 1} parsing error: {e}")
        
        return panels_data
    
    def _parse_copy_table(self, table, doc: Document) -> Dict[str, Dict[str, Dict[str, str]]]:
        """
        Parse a single table from the copy document.
        
        Detects panel headers, language columns, and field data.
        Also handles strikethrough detection at the run level.
        
        Args:
            table: python-docx Table object
            doc: Parent document (for style access)
            
        Returns:
            Nested dict structure for this table
        """
        all_panels_data = {}
        current_panel = None
        lang_columns = {}  # col_index -> language_code
        
        for row in table.rows:
            cells = []
            
            # Extract cell text with strikethrough detection
            for cell in row.cells:
                cell_text = ""
                has_strikethrough = False
                
                for para in cell.paragraphs:
                    for run in para.runs:
                        # Check for strikethrough
                        if run.font.strike:
                            has_strikethrough = True
                            # Still collect text but mark it
                            cell_text += f"[STRIKE]{run.text}[/STRIKE]"
                        else:
                            cell_text += run.text
                
                cells.append(cell_text.strip())
            
            if not cells or all(not c for c in cells):
                continue
            
            # Detect panel header rows
            first_cell_lower = cells[0].lower() if cells[0] else ""
            if 'front of artwork' in first_cell_lower:
                current_panel = 'Front Panel'
                all_panels_data[current_panel] = {}
                lang_columns = {}
                continue
            elif 'back of artwork' in first_cell_lower:
                current_panel = 'Back Panel'
                all_panels_data[current_panel] = {}
                lang_columns = {}
                continue
            elif 'side of artwork' in first_cell_lower:
                current_panel = 'Side Panel'
                all_panels_data[current_panel] = {}
                lang_columns = {}
                continue
            
            # Detect language header row
            if current_panel and not lang_columns:
                potential_langs = {}
                for i, cell in enumerate(cells):
                    cell_upper = cell.upper().strip()
                    if cell_upper == 'ENGLISH':
                        potential_langs[i] = 'EN'
                    elif cell_upper == 'FRENCH':
                        potential_langs[i] = 'FR'
                    elif cell_upper == 'SPANISH':
                        potential_langs[i] = 'ES'
                    elif cell_upper == 'GERMAN':
                        potential_langs[i] = 'DE'
                    elif cell_upper == 'DUTCH':
                        potential_langs[i] = 'NL'
                    elif cell_upper == 'ITALIAN':
                        potential_langs[i] = 'IT'
                    elif cell_upper == 'PORTUGUESE':
                        potential_langs[i] = 'PT'
                    elif cell_upper == 'POLISH':
                        potential_langs[i] = 'PL'
                    elif 'SCAND' in cell_upper or 'DANISH' in cell_upper:
                        potential_langs[i] = 'DA'
                    elif cell_upper == 'FINNISH':
                        potential_langs[i] = 'FI'
                    elif cell_upper == 'RUSSIAN':
                        potential_langs[i] = 'RU'
                
                if potential_langs:
                    lang_columns = potential_langs
                    for lang in lang_columns.values():
                        if lang not in all_panels_data[current_panel]:
                            all_panels_data[current_panel][lang] = {}
                    continue
            
            # Data row processing
            if current_panel and lang_columns and cells[0]:
                field_name = cells[0]
                
                # Skip if this looks like another header
                if field_name.lower() in ['front of artwork', 'back of artwork', 'side of artwork']:
                    continue
                
                # Extract values for each language
                for col_idx, language in lang_columns.items():
                    if col_idx < len(cells):
                        value = cells[col_idx]
                        if value:
                            all_panels_data[current_panel][language][field_name] = value
        
        return all_panels_data
    
    def _build_field_list(
        self, 
        panels_data: Dict[str, Dict[str, Dict[str, str]]],
        doc: Document
    ) -> List[CopyField]:
        """
        Build flat list of CopyField objects from nested panel data.
        
        Handles:
        - Field splitting for multi-paragraph fields
        - Strikethrough detection
        - Legacy text (` -> `) separation
        - Pack claim numbering
        
        Args:
            panels_data: Nested dict from table extraction
            doc: Original document for strikethrough re-check
            
        Returns:
            Flat list of CopyField objects
        """
        fields = []
        
        for panel, languages in panels_data.items():
            for language, field_data in languages.items():
                for field_name, text in field_data.items():
                    # Handle strikethrough markers
                    is_strikethrough = '[STRIKE]' in text and '[/STRIKE]' in text
                    
                    # Clean strikethrough markers for processing
                    clean_text = re.sub(r'\[STRIKE\].*?\[/STRIKE\]', '', text)
                    clean_text = clean_text.strip()
                    
                    # Handle legacy text (` -> ` separator)
                    if ' -> ' in clean_text:
                        parts = clean_text.split(' -> ')
                        
                        # First part is legacy (before arrow)
                        if parts[0].strip():
                            fields.append(CopyField(
                                field_name=field_name,
                                panel=panel,
                                language=language,
                                text=parts[0].strip(),
                                is_strikethrough=is_strikethrough,
                                is_legacy=True
                            ))
                        
                        # Last part is current (after arrow)
                        if len(parts) > 1 and parts[-1].strip():
                            clean_text = parts[-1].strip()
                    
                    # Check if this field should be split
                    if self._should_split_field(field_name):
                        split_fields = self._split_field(
                            field_name, panel, language, clean_text
                        )
                        fields.extend(split_fields)
                    else:
                        fields.append(CopyField(
                            field_name=field_name,
                            panel=panel,
                            language=language,
                            text=clean_text,
                            is_strikethrough=is_strikethrough,
                            is_legacy=False
                        ))
        
        return fields
    
    def _should_split_field(self, field_name: str) -> bool:
        """
        Check if a field should be split into sub-fields.
        
        Args:
            field_name: Name of the field
            
        Returns:
            True if field should be split
        """
        for split_field in config.PARAGRAPH_SPLIT_FIELDS:
            if split_field.lower() in field_name.lower():
                return True
        return False
    
    def _split_field(
        self,
        field_name: str,
        panel: str,
        language: str,
        text: str
    ) -> List[CopyField]:
        """
        Split a field into sub-fields by paragraph.
        
        Args:
            field_name: Original field name
            panel: Panel location
            language: Language code
            text: Text to split
            
        Returns:
            List of CopyField objects for each sub-field
        """
        fields = []
        
        # Split by newlines (single or double)
        paragraphs = re.split(r'\n+', text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        
        if not paragraphs:
            return fields
        
        # Get template for this field type
        template = None
        for key, tmpl in config.SUBFIELD_TEMPLATES.items():
            if key.lower() in field_name.lower():
                template = tmpl
                break
        
        # Create sub-fields
        for i, para in enumerate(paragraphs):
            if template:
                if '{}' in template[0]:
                    # Numbered template (e.g., "Hero Ingredient 1")
                    subfield_name = template[0].format(i + 1)
                elif i < len(template):
                    # Named template (e.g., "Marketing Copy", "Scent Line")
                    subfield_name = template[i]
                else:
                    subfield_name = f"{field_name} - Part {i + 1}"
            else:
                subfield_name = f"{field_name} - Part {i + 1}"
            
            fields.append(CopyField(
                field_name=subfield_name,
                panel=panel,
                language=language,
                text=para,
                parent_field=field_name,
                subfield_index=i
            ))
        
        return fields




# =============================================================================
# PDF EXTRACTOR
# =============================================================================

class PDFExtractor:
    """
    Extracts text from PDF artwork files.
    
    Uses PyMuPDF (fitz) for reliable text extraction with
    font metadata. Handles outlined text detection and
    provides extraction confidence metrics.
    """
    
    def __init__(self):
        """Initialize the PDF extractor."""
        self.warnings = []
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """
        Extract all text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            ExtractionResult with all extracted text runs
        """
        self.warnings = []
        
        if not PYMUPDF_AVAILABLE:
            self.warnings.append("PyMuPDF (fitz) not available")
            return ExtractionResult(
                file_path=str(file_path),
                file_type="pdf",
                extraction_method=ExtractionMethod.FAILED,
                confidence=0.0,
                warnings=self.warnings
            )
        
        try:
            doc = fitz.open(file_path)
            page_count = len(doc)  # Get this BEFORE closing
            logger.info(f"Opened PDF: {file_path.name} ({len(doc)} pages)")
            
            all_text_runs = []
            pages_with_text = 0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_runs = self._extract_page_text(page, page_num + 1)
                
                if page_runs:
                    all_text_runs.extend(page_runs)
                    pages_with_text += 1
                else:
                    self.warnings.append(f"Page {page_num + 1}: No extractable text found")
            
            # Calculate confidence BEFORE closing
            if len(doc) == 0:
                confidence = 0.0
                method = ExtractionMethod.FAILED
            elif pages_with_text == 0:
                confidence = 0.0
                method = ExtractionMethod.FAILED
                self.warnings.append("No text extracted - PDF may have outlined fonts")
            else:
                confidence = pages_with_text / len(doc)
                method = ExtractionMethod.LIVE_TEXT

            # NOW close
            doc.close()
            
            # Apply exclusion filter
            filtered_runs = ExclusionChecker.filter_text_runs(all_text_runs)
            
            return ExtractionResult(
                file_path=str(file_path),
                file_type="pdf",
                extraction_method=method,
                confidence=confidence,
                text_runs=filtered_runs,
                pages_processed=len(doc),
                pages_with_text=pages_with_text,
                warnings=self.warnings,
                metadata={
                    'total_runs_before_filter': len(all_text_runs),
                    'total_runs_after_filter': len(filtered_runs),
                    'page_count': len(doc)
                }
            )
            
        except Exception as e:
            self.warnings.append(f"PDF extraction error: {str(e)}")
            logger.error(f"Failed to extract PDF: {e}")
            return ExtractionResult(
                file_path=str(file_path),
                file_type="pdf",
                extraction_method=ExtractionMethod.FAILED,
                confidence=0.0,
                warnings=self.warnings
            )
    
    def _extract_page_text(self, page, page_num: int) -> List[TextRun]:
        """
        Extract text runs from a single page.
        
        Args:
            page: PyMuPDF page object
            page_num: 1-based page number
            
        Returns:
            List of TextRun objects from this page
        """
        text_runs = []
        
        try:
            # Extract with full metadata
            blocks = page.get_text("dict", flags=11)["blocks"]
            
            for block in blocks:
                # Only process text blocks
                if block.get("type") != 0:
                    continue
                
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        
                        if not text:
                            continue
                        
                        bbox = span.get("bbox")
                        font_name = span.get("font", "Unknown")
                        font_size = span.get("size", 0.0)
                        
                        # Detect rotation from line direction
                        rotation = self._detect_rotation(line)
                        
                        text_run = TextRun(
                            text=text,
                            page_number=page_num,
                            bbox=tuple(bbox) if bbox else None,
                            font_name=font_name,
                            font_size=round(font_size, 2),
                            extraction_method=ExtractionMethod.LIVE_TEXT,
                            confidence=1.0,
                            rotation=rotation
                        )
                        text_runs.append(text_run)
            
            # Fallback: try simple text extraction if no runs found
            if not text_runs:
                simple_text = page.get_text().strip()
                if simple_text and len(simple_text) >= 10:
                    text_run = TextRun(
                        text=simple_text,
                        page_number=page_num,
                        extraction_method=ExtractionMethod.LIVE_TEXT,
                        confidence=0.7
                    )
                    text_runs.append(text_run)
                    
        except Exception as e:
            logger.warning(f"Page {page_num} extraction error: {e}")
        
        return text_runs
    
    def _detect_rotation(self, line: Dict) -> int:
        """
        Detect text rotation from line direction vector.
        
        Args:
            line: PyMuPDF line dict with 'dir' field
            
        Returns:
            Rotation in degrees (0, 90, 180, 270)
        """
        dir_vec = line.get("dir", (1, 0))
        
        if len(dir_vec) < 2:
            return 0
        
        dx, dy = dir_vec
        
        # Determine rotation based on direction
        if abs(dx - 1) < 0.1 and abs(dy) < 0.1:
            return 0  # Normal horizontal
        elif abs(dx) < 0.1 and abs(dy - 1) < 0.1:
            return 90  # Vertical, bottom to top
        elif abs(dx + 1) < 0.1 and abs(dy) < 0.1:
            return 180  # Upside down
        elif abs(dx) < 0.1 and abs(dy + 1) < 0.1:
            return 270  # Vertical, top to bottom
        
        return 0


# =============================================================================
# AI FILE EXTRACTOR
# =============================================================================

class AIExtractor:
    """
    Extracts text from Adobe Illustrator (.ai) files.
    
    AI files with PDF compatibility can be read as PDFs.
    Falls back to failed status if PDF compatibility is not present.
    """
    
    def __init__(self):
        """Initialize the AI extractor."""
        self.pdf_extractor = PDFExtractor()
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """
        Extract text from an AI file.
        
        Args:
            file_path: Path to the .ai file
            
        Returns:
            ExtractionResult (may be via PDF compatibility)
        """
        # Try PDF compatibility extraction
        result = self.pdf_extractor.extract(file_path)
        
        if result.confidence > 0:
            result.extraction_method = ExtractionMethod.AI_PDF_COMPATIBLE
            result.file_type = "ai"
            result.warnings.insert(0, "AI file extracted via PDF compatibility")
            logger.info(f"AI file extracted via PDF compatibility: {file_path.name}")
        else:
            result.warnings.append("AI file may not have PDF compatibility enabled")
            logger.warning(f"AI extraction failed for: {file_path.name}")
        
        return result


# =============================================================================
# FONT EXTRACTOR
# =============================================================================

class FontExtractor:
    """
    Extracts font size metadata from artwork.
    
    Used for Section 3E to verify minimum font size
    requirements for regulatory compliance.
    """
    
    def extract_from_pdf(self, file_path: Path) -> List[FontMeasurement]:
        """
        Extract all font measurements from a PDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of FontMeasurement objects
        """
        if not PYMUPDF_AVAILABLE:
            logger.warning("PyMuPDF not available for font extraction")
            return []
        
        measurements = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = page.get_text("dict", flags=11)["blocks"]
                
                for block in blocks:
                    if block.get("type") != 0:
                        continue
                    
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if not text:
                                continue
                            
                            # Skip excluded text
                            should_exclude, _ = ExclusionChecker.should_exclude(text)
                            if should_exclude:
                                continue
                            
                            font_size = span.get("size", 0)
                            if font_size > 0:
                                measurements.append(FontMeasurement(
                                    text=text[:50],
                                    page_number=page_num + 1,
                                    font_name=span.get("font", "Unknown"),
                                    font_size_pt=round(font_size, 2),
                                    bbox=tuple(span.get("bbox", []))
                                ))
            
            doc.close()
            logger.info(f"Extracted {len(measurements)} font measurements")
            
        except Exception as e:
            logger.error(f"Font extraction error: {e}")
        
        return measurements
    
    def get_smallest_font(self, measurements: List[FontMeasurement]) -> Optional[FontMeasurement]:
        """
        Find the smallest font in measurements.
        
        Args:
            measurements: List of font measurements
            
        Returns:
            FontMeasurement with smallest size, or None
        """
        if not measurements:
            return None
        return min(measurements, key=lambda m: m.font_size_pt)


# =============================================================================
# BARCODE SCANNER
# =============================================================================

class BarcodeScanner:
    """
    Scans and validates barcodes in artwork.
    
    Section 3F - Uses pyzbar to decode barcodes from
    rendered PDF pages and validates check digits.
    """
    
    @staticmethod
    def scan_from_pdf(file_path: Path) -> List[BarcodeResult]:
        """
        Scan for barcodes in a PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of BarcodeResult objects
        """
        results = []
        
        if not PYMUPDF_AVAILABLE:
            return [BarcodeResult(
                symbology="Unknown",
                decoded_digits=None,
                printed_digits=None,
                digits_match=False,
                check_digit_valid=False,
                scan_successful=False,
                quality_notes="PyMuPDF not available for barcode scanning",
                status_code=StatusCode.TBD
            )]
        
        if not PYZBAR_AVAILABLE:
            return [BarcodeResult(
                symbology="Unknown",
                decoded_digits=None,
                printed_digits=None,
                digits_match=False,
                check_digit_valid=False,
                scan_successful=False,
                quality_notes="pyzbar not available - manual verification required",
                status_code=StatusCode.TBD
            )]
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Render at high DPI for barcode detection
                mat = fitz.Matrix(config.BARCODE_RENDER_DPI, config.BARCODE_RENDER_DPI)
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to PIL Image
                img_data = pix.tobytes("png")
                img = Image.open(BytesIO(img_data))
                
                # Decode barcodes
                decoded = decode_barcode(img)
                
                for barcode in decoded:
                    barcode_type = barcode.type
                    barcode_data = barcode.data.decode('utf-8')
                    
                    # Validate check digit
                    is_valid = BarcodeScanner._validate_check_digit(barcode_data, barcode_type)
                    
                    results.append(BarcodeResult(
                        symbology=barcode_type,
                        decoded_digits=barcode_data,
                        printed_digits=barcode_data,  # Would need OCR for printed digits
                        digits_match=True,  # Assuming decoded = printed for now
                        check_digit_valid=is_valid,
                        scan_successful=True,
                        quality_notes=f"Successfully decoded on page {page_num + 1}",
                        status_code=StatusCode.OK if is_valid else StatusCode.FAIL
                    ))
            
            doc.close()
            
            if not results:
                results.append(BarcodeResult(
                    symbology="N/A",
                    decoded_digits=None,
                    printed_digits=None,
                    digits_match=False,
                    check_digit_valid=False,
                    scan_successful=False,
                    quality_notes="No barcode detected in artwork",
                    status_code=StatusCode.FYI
                ))
            
        except Exception as e:
            logger.error(f"Barcode scanning error: {e}")
            results.append(BarcodeResult(
                symbology="Unknown",
                decoded_digits=None,
                printed_digits=None,
                digits_match=False,
                check_digit_valid=False,
                scan_successful=False,
                quality_notes=f"Scan error: {str(e)}",
                status_code=StatusCode.TBD
            ))
        
        return results
    
    @staticmethod
    def _validate_check_digit(digits: str, barcode_type: str) -> bool:
        """
        Validate barcode check digit.
        
        Args:
            digits: Barcode digit string
            barcode_type: Type of barcode (EAN13, UPCA, etc.)
            
        Returns:
            True if check digit is valid
        """
        digits = digits.replace(" ", "").replace("-", "")
        
        if barcode_type in ["EAN13", "EAN-13"] and len(digits) == 13:
            return BarcodeScanner._validate_ean13(digits)
        elif barcode_type in ["UPCA", "UPC-A", "UPC_A"] and len(digits) == 12:
            return BarcodeScanner._validate_upc_a(digits)
        elif barcode_type in ["EAN8", "EAN-8"] and len(digits) == 8:
            return BarcodeScanner._validate_ean8(digits)
        
        return True  # Unknown type, assume valid
    
    @staticmethod
    def _validate_ean13(digits: str) -> bool:
        """Validate EAN-13 check digit."""
        if len(digits) != 13 or not digits.isdigit():
            return False
        
        odds = sum(int(digits[i]) for i in range(0, 12, 2))
        evens = sum(int(digits[i]) for i in range(1, 12, 2))
        check = (10 - ((odds + evens * 3) % 10)) % 10
        
        return check == int(digits[12])
    
    @staticmethod
    def _validate_upc_a(digits: str) -> bool:
        """Validate UPC-A check digit."""
        if len(digits) != 12 or not digits.isdigit():
            return False
        
        odds = sum(int(digits[i]) for i in range(0, 11, 2))
        evens = sum(int(digits[i]) for i in range(1, 11, 2))
        check = (10 - ((odds * 3 + evens) % 10)) % 10
        
        return check == int(digits[11])
    
    @staticmethod
    def _validate_ean8(digits: str) -> bool:
        """Validate EAN-8 check digit."""
        if len(digits) != 8 or not digits.isdigit():
            return False
        
        odds = sum(int(digits[i]) for i in range(0, 7, 2))
        evens = sum(int(digits[i]) for i in range(1, 7, 2))
        check = (10 - ((odds * 3 + evens) % 10)) % 10
        
        return check == int(digits[7])


    #             [rel_x0, rel_y0, rel_x1, rel_y1],
    #             outline=color,
    #             width=3
    #         )
    #         
    #         # Add issue ID label
    #         try:
    #             font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 16)
    #         except:
    #             font = ImageFont.load_default()
    #         
    #         draw.text((5, 5), issue_id, fill=color, font=font)
    #         
    #         img.save(output_path, "PNG")
    #         doc.close()
    #         
    #         return True
    #         
    #     except Exception as e:
    #         logger.error(f"Crop snapshot error: {e}")
    #         return False
    # 
    # END PLACEHOLDER
    # =========================================================================
    
    @classmethod
    def generate_snapshots_for_findings(
        cls,
        pdf_path: Path,
        findings: List[MatchFinding],
        output_dir: Path
    ) -> List[str]:
        """
        Generate snapshots for all findings with issues.
        
        Args:
            pdf_path: Path to artwork PDF
            findings: List of match findings
            output_dir: Directory to save snapshots
            
        Returns:
            List of generated snapshot file paths
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        generated_files = []
        
        # Group annotations by page
        annotations_by_page: Dict[int, List[SnapshotAnnotation]] = {}
        
        for finding in findings:
            if finding.status_code == StatusCode.OK:
                continue  # Only annotate issues
            
            if not finding.bbox:
                continue
            
            page_num = 1  # Default to page 1 if not specified
            
            annotation = SnapshotAnnotation(
                issue_id=finding.issue_id or "?",
                bbox=finding.bbox,
                status=finding.status_code,
                page_number=page_num
            )
            
            if page_num not in annotations_by_page:
                annotations_by_page[page_num] = []
            
            if len(annotations_by_page[page_num]) < config.SNAPSHOT_MAX_PER_PAGE:
                annotations_by_page[page_num].append(annotation)
        
        # Generate one snapshot per page with annotations
        for page_num, annotations in annotations_by_page.items():
            output_path = output_dir / f"snapshot_page_{page_num}.png"
            
            success = cls.generate_annotated_page(
                pdf_path,
                page_num,
                annotations,
                output_path
            )
            
            if success:
                generated_files.append(str(output_path))
        
        return generated_files

# =============================================================================
# ARTWORK MATCHER
# =============================================================================

class ArtworkMatcher:
    """
    Matches copy document fields against extracted artwork text.
    
    Implements strict deterministic matching with anti-hallucination safeguards.
    Uses sliding window approach to match multi-span fields.
    
    Attributes:
        config: Configuration instance
        normalizer: Text normalization utility
        zoom_detector: Zoom trigger detection utility
    """
    
    def __init__(self, config: Config):
        """
        Initialize the artwork matcher.
        
        Args:
            config: Configuration instance
        """
        self.config = config
        self.normalizer = TextNormalizer(config)
        self.zoom_detector = ZoomTriggerDetector(config)
        logger.info("ArtworkMatcher initialized")
    
    def match_fields(
        self,
        copy_fields: List[CopyField],
        text_runs: List[TextRun],
        extraction_method: ExtractionMethod
    ) -> List[MatchFinding]:
        """
        Match copy fields to artwork text with strict evidence requirements.
        
        ANTI-HALLUCINATION INVARIANTS:
        1. If extraction_method == FAILED → NO matches generated
        2. If no matching text_run found → MatchType.MISSING_IN_ARTWORK
        3. EXACT_MATCH only if normalized strings are identical
        4. All confidence scores must be grounded in actual fuzzy match scores
        
        Args:
            copy_fields: Fields from copy document
            text_runs: Text extracted from artwork
            extraction_method: How text was extracted
            
        Returns:
            List of match findings with evidence
        """
        findings = []
        
        # FAIL-FAST: If extraction failed, cannot match
        if extraction_method == ExtractionMethod.FAILED:
            logger.warning("Extraction failed - cannot generate reliable matches")
            for field in copy_fields:
                findings.append(MatchFinding(
                    copy_field=field,
                    match_type=MatchType.REQUIRES_VERIFICATION,
                    status=StatusCode.TBD,
                    artwork_value=None,
                    matched_runs=[],
                    fuzzy_score=0.0,
                    zoom_triggers=[],
                    notes="Extraction failed - manual verification required"
                ))
            return findings
        
        # Build lookup for fast exact matching
        artwork_text_lookup = self._build_text_lookup(text_runs)
        
        # Match each copy field
        logger.info(f"Matching {len(copy_fields)} copy fields against {len(text_runs)} text runs")
        
        for field in copy_fields:
            finding = self._match_single_field(field, text_runs, artwork_text_lookup)
            findings.append(finding)
        
        # Log summary
        exact_matches = sum(1 for f in findings if f.match_type == MatchType.EXACT_MATCH)
        near_matches = sum(1 for f in findings if f.match_type == MatchType.NEAR_MATCH)
        mismatches = sum(1 for f in findings if f.match_type == MatchType.MISMATCH)
        missing = sum(1 for f in findings if f.match_type == MatchType.MISSING_IN_ARTWORK)
        
        logger.info(f"Match results: {exact_matches} exact, {near_matches} near, "
                   f"{mismatches} mismatch, {missing} missing")
        
        return findings
    
    def _match_single_field(
        self,
        field: CopyField,
        text_runs: List[TextRun],
        artwork_lookup: Dict[str, List[TextRun]]
    ) -> MatchFinding:
        """
        Match a single copy field with strict evidence requirements.
        
        Args:
            field: Copy field to match
            text_runs: All text runs from artwork
            artwork_lookup: Normalized text → runs mapping
            
        Returns:
            MatchFinding with actual evidence or MISSING status
        """
        normalized_copy = self.normalizer.normalize(field.text)
        
        # Check zoom triggers
        zoom_triggers = self.zoom_detector.check_triggers(field.text)
        requires_visual = bool(zoom_triggers)
        
        # Try exact match first
        if normalized_copy in artwork_lookup:
            matched_runs = artwork_lookup[normalized_copy]
            return MatchFinding(
                copy_field=field,
                match_type=MatchType.EXACT_MATCH,
                status=StatusCode.TBD if requires_visual else StatusCode.OK,
                artwork_value=field.text,  # Use copy value for exact match
                matched_runs=matched_runs,
                fuzzy_score=100.0,
                zoom_triggers=zoom_triggers,
                notes=self._generate_match_notes(True, requires_visual, zoom_triggers)
            )
        
        # Try fuzzy matching with sliding window for multi-span matches
        best_match = self._fuzzy_match_with_sliding_window(normalized_copy, text_runs)
        
        if best_match:
            match_type, fuzzy_score, matched_text, matched_runs = best_match
            
            if fuzzy_score >= self.config.EXACT_MATCH_THRESHOLD:
                # Fuzzy matcher found 100% match
                return MatchFinding(
                    copy_field=field,
                    match_type=MatchType.EXACT_MATCH,
                    status=StatusCode.TBD if requires_visual else StatusCode.OK,
                    artwork_value=matched_text,
                    matched_runs=matched_runs,
                    fuzzy_score=fuzzy_score,
                    zoom_triggers=zoom_triggers,
                    notes=self._generate_match_notes(True, requires_visual, zoom_triggers)
                )
            elif fuzzy_score >= self.config.NEAR_MATCH_THRESHOLD:
                # Near match
                return MatchFinding(
                    copy_field=field,
                    match_type=MatchType.NEAR_MATCH,
                    status=StatusCode.TBD if requires_visual else StatusCode.ATTN,
                    artwork_value=matched_text,
                    matched_runs=matched_runs,
                    fuzzy_score=fuzzy_score,
                    zoom_triggers=zoom_triggers,
                    notes=f"Near match ({fuzzy_score:.1f}%) - verify differences" +
                          (f" | Zoom: {', '.join(zoom_triggers)}" if zoom_triggers else "")
                )
            else:
                # Mismatch
                return MatchFinding(
                    copy_field=field,
                    match_type=MatchType.MISMATCH,
                    status=StatusCode.FAIL,
                    artwork_value=matched_text,
                    matched_runs=matched_runs,
                    fuzzy_score=fuzzy_score,
                    zoom_triggers=zoom_triggers,
                    notes=f"Mismatch detected ({fuzzy_score:.1f}%)" +
                          (f" | Zoom: {', '.join(zoom_triggers)}" if zoom_triggers else "")
                )
        
        # NO MATCH FOUND - This is NOT a hallucination; it's evidence of absence
        return MatchFinding(
            copy_field=field,
            match_type=MatchType.MISSING_IN_ARTWORK,
            status=StatusCode.FAIL,
            artwork_value=None,
            matched_runs=[],
            fuzzy_score=0.0,
            zoom_triggers=zoom_triggers,
            notes="🚨 NOT FOUND in artwork - requires visual confirmation" +
                  (f" | Zoom: {', '.join(zoom_triggers)}" if zoom_triggers else "")
        )
    
    def _build_text_lookup(self, text_runs: List[TextRun]) -> Dict[str, List[TextRun]]:
        """
        Build normalized text → runs mapping for fast exact matching.
        
        Args:
            text_runs: All text runs from artwork
            
        Returns:
            Dictionary mapping normalized text to list of runs
        """
        lookup = {}
        for run in text_runs:
            normalized = self.normalizer.normalize(run.text)
            if normalized:  # Skip empty normalized text
                if normalized not in lookup:
                    lookup[normalized] = []
                lookup[normalized].append(run)
        return lookup
    
    def _fuzzy_match_with_sliding_window(
        self,
        normalized_copy: str,
        text_runs: List[TextRun]
    ) -> Optional[Tuple[MatchType, float, str, List[TextRun]]]:
        """
        Attempt fuzzy matching using sliding window for multi-span fields.
        
        This handles the case where a copy field spans multiple TextRuns
        (e.g., a multi-paragraph ingredient list).
        
        Args:
            normalized_copy: Normalized copy field text
            text_runs: All text runs from artwork
            
        Returns:
            (match_type, fuzzy_score, matched_text, matched_runs) or None
        """
        best_score = 0.0
        best_match = None
        
        # Try individual runs first
        for run in text_runs:
            normalized_artwork = self.normalizer.normalize(run.text)
            if not normalized_artwork:
                continue
            
            score = SequenceMatcher(None, normalized_copy, normalized_artwork).ratio() * 100
            
            if score > best_score:
                best_score = score
                best_match = (normalized_artwork, [run])
        
        # Try sliding window of concatenated runs (for multi-span matches)
        # Window sizes from 2 to 5 runs
        for window_size in range(2, min(6, len(text_runs) + 1)):
            for i in range(len(text_runs) - window_size + 1):
                window_runs = text_runs[i:i+window_size]
                
                # Concatenate normalized text from window
                concatenated = ' '.join(
                    self.normalizer.normalize(r.text) 
                    for r in window_runs
                    if self.normalizer.normalize(r.text)
                )
                
                if not concatenated:
                    continue
                
                score = SequenceMatcher(None, normalized_copy, concatenated).ratio() * 100
                
                if score > best_score:
                    best_score = score
                    best_match = (concatenated, window_runs)
        
        # Only return if score meets minimum threshold
        if best_match and best_score >= self.config.MISMATCH_THRESHOLD:
            matched_text, matched_runs = best_match
            
            # Determine match type based on score
            if best_score >= self.config.EXACT_MATCH_THRESHOLD:
                match_type = MatchType.EXACT_MATCH
            elif best_score >= self.config.NEAR_MATCH_THRESHOLD:
                match_type = MatchType.NEAR_MATCH
            else:
                match_type = MatchType.MISMATCH
            
            return (match_type, best_score, matched_text, matched_runs)
        
        return None
    
    def _generate_match_notes(
        self,
        is_exact: bool,
        requires_visual: bool,
        zoom_triggers: List[str]
    ) -> str:
        """
        Generate appropriate notes for a match finding.
        
        Args:
            is_exact: Whether this is an exact match
            requires_visual: Whether visual verification is required
            zoom_triggers: List of zoom trigger reasons
            
        Returns:
            Formatted notes string
        """
        if is_exact and not requires_visual:
            return "✅ Exact match"
        elif is_exact and requires_visual:
            return f"🔍 Exact match - VERIFY VISUALLY | Zoom: {', '.join(zoom_triggers)}"
        else:
            return ""


# =============================================================================
# CONVERSION CHECKER
# =============================================================================

class ConversionChecker:
    """
    Verifies volume conversions between mL and fl oz.
    
    Checks that declared volumes follow correct conversion math
    and fall within acceptable tolerance ranges.
    
    Attributes:
        config: Configuration instance
        ML_TO_FL_OZ: Standard conversion factor (0.033814)
        TOLERANCE: Acceptable difference in fl oz (0.10)
    """
    
    ML_TO_FL_OZ = 0.033814  # Standard mL to fl oz conversion
    TOLERANCE = 0.10        # ±0.10 fl oz acceptable
    
    def __init__(self, config: Config):
        """
        Initialize the conversion checker.
        
        Args:
            config: Configuration instance
        """
        self.config = config
        logger.info("ConversionChecker initialized")
    
    def check_conversions(
        self,
        copy_fields: List[CopyField],
        artwork_text_runs: List[TextRun]
    ) -> List[ConversionCheck]:
        """
        Find and verify all volume conversions in the artwork.
        
        Args:
            copy_fields: Fields from copy document (to find fill weight field)
            artwork_text_runs: Extracted artwork text
            
        Returns:
            List of conversion check results
        """
        results = []
        
        # Find fill weight in copy doc
        fill_weight_field = self._find_fill_weight(copy_fields)
        if not fill_weight_field:
            logger.warning("No fill weight field found in copy document")
            return results
        
        logger.info(f"Found fill weight field: {fill_weight_field.text}")
        
        # Extract volumes from fill weight
        ml_value, fl_oz_value = self._extract_volumes(fill_weight_field.text)
        
        if ml_value and fl_oz_value:
            # Calculate expected conversion
            calculated_fl_oz = ml_value * self.ML_TO_FL_OZ
            difference = abs(fl_oz_value - calculated_fl_oz)
            
            # Check if within tolerance
            passed = difference <= self.TOLERANCE
            
            result = ConversionCheck(
                field_name="Fill Weight",
                ml_value=ml_value,
                fl_oz_value=fl_oz_value,
                calculated_fl_oz=calculated_fl_oz,
                difference=difference,
                tolerance=self.TOLERANCE,
                passed=passed,
                status=StatusCode.OK if passed else StatusCode.FAIL,
                notes=f"Declared: {fl_oz_value} fl oz | Calculated: {calculated_fl_oz:.2f} fl oz | Diff: {difference:.2f}"
            )
            
            logger.info(f"Conversion check: {ml_value} mL → {fl_oz_value} fl oz "
                       f"(calculated {calculated_fl_oz:.2f}, diff {difference:.2f}) - "
                       f"{'PASS' if passed else 'FAIL'}")
            
            results.append(result)
        else:
            logger.warning(f"Could not extract volumes from fill weight: {fill_weight_field.text}")
        
        return results
    
    def _find_fill_weight(self, copy_fields: List[CopyField]) -> Optional[CopyField]:
        """
        Find the fill weight field in copy document.
        
        Args:
            copy_fields: All fields from copy document
            
        Returns:
            Fill weight field or None
        """
        for field in copy_fields:
            if 'fill weight' in field.field_name.lower():
                return field
        return None
    
    def _extract_volumes(self, text: str) -> Tuple[Optional[float], Optional[float]]:
        """
        Extract mL and fl oz values from fill weight text.
        
        Examples:
            "250 ML / 8.5 US FL. OZ." → (250.0, 8.5)
            "400ML / 13.5 FL OZ" → (400.0, 13.5)
            "30 ml / 1.0 fl oz" → (30.0, 1.0)
        
        Args:
            text: Fill weight text to parse
            
        Returns:
            (ml_value, fl_oz_value) or (None, None)
        """
        # Pattern for mL (handles various formats)
        ml_pattern = r'(\d+(?:\.\d+)?)\s*(?:ML|ml|mL|Ml)\b'
        
        # Pattern for fl oz (handles US, UK variations)
        fl_oz_pattern = r'(\d+(?:\.\d+)?)\s*(?:US\s*)?(?:FL\.?\s*OZ\.?|fl\.?\s*oz\.?|Fl\.?\s*Oz\.?)\b'
        
        ml_match = re.search(ml_pattern, text, re.IGNORECASE)
        fl_oz_match = re.search(fl_oz_pattern, text, re.IGNORECASE)
        
        ml_value = float(ml_match.group(1)) if ml_match else None
        fl_oz_value = float(fl_oz_match.group(1)) if fl_oz_match else None
        
        return ml_value, fl_oz_value


# =============================================================================
# SNAPSHOT GENERATOR
# =============================================================================

class SnapshotGenerator:
    """
    Generates annotated full-page snapshots of artwork with issue highlights.
    
    Creates visual reference images showing where issues are located
    on the artwork, with bounding boxes and labels color-coded by severity.
    
    Attributes:
        config: Configuration instance
    """
    
    def __init__(self, config: Config):
        """
        Initialize the snapshot generator.
        
        Args:
            config: Configuration instance
        """
        self.config = config
        logger.info("SnapshotGenerator initialized")
    
    def generate_snapshots(
        self,
        pdf_path: Path,
        findings: List[MatchFinding],
        output_dir: Path
    ) -> List[Path]:
        """
        Generate annotated snapshot images for visual reference.
        
        Args:
            pdf_path: Path to artwork PDF
            findings: Match findings with matched_runs containing bbox data
            output_dir: Where to save snapshot images
            
        Returns:
            List of paths to generated snapshot images
        """
        if not PIL_AVAILABLE:
            logger.warning("PIL not available - cannot generate snapshots")
            return []
        
        if not PYMUPDF_AVAILABLE:
            logger.warning("PyMuPDF not available - cannot generate snapshots")
            return []
        
        if not self.config.SNAPSHOT_ENABLED:
            logger.info("Snapshots disabled in configuration")
            return []
        
        snapshot_paths = []
        
        try:
            # Open PDF
            doc = fitz.open(str(pdf_path))
            
            # Group findings by page
            findings_by_page = self._group_findings_by_page(findings)
            
            logger.info(f"Generating snapshots for {len(findings_by_page)} pages")
            
            for page_num, page_findings in findings_by_page.items():
                if page_num >= len(doc):
                    logger.warning(f"Page {page_num} out of range (doc has {len(doc)} pages)")
                    continue
                
                page = doc[page_num]
                
                # Render page to image
                pix = page.get_pixmap(dpi=self.config.SNAPSHOT_DPI)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Draw annotations
                img = self._draw_annotations(img, page_findings, page, pix)
                
                # Save
                output_dir.mkdir(parents=True, exist_ok=True)
                snapshot_path = output_dir / f"snapshot_page_{page_num + 1}.png"
                img.save(snapshot_path)
                snapshot_paths.append(snapshot_path)
                
                logger.info(f"Generated snapshot: {snapshot_path}")
            
            doc.close()
            
        except Exception as e:
            logger.error(f"Error generating snapshots: {e}")
        
        return snapshot_paths
    
    def _group_findings_by_page(
        self,
        findings: List[MatchFinding]
    ) -> Dict[int, List[MatchFinding]]:
        """
        Group findings by page number for efficient rendering.
        
        Args:
            findings: All match findings
            
        Returns:
            Dictionary mapping page number to list of findings on that page
        """
        by_page = {}
        
        for finding in findings:
            if not finding.matched_runs:
                continue
            
            # Use the page number from the first matched run
            for run in finding.matched_runs:
                if run.bbox:
                    page_num = run.page_num
                    if page_num not in by_page:
                        by_page[page_num] = []
                    by_page[page_num].append(finding)
                    break  # Only need one page per finding
        
        return by_page
    
    def _draw_annotations(
        self,
        img: Image.Image,
        findings: List[MatchFinding],
        page: fitz.Page,
        pix: fitz.Pixmap
    ) -> Image.Image:
        """
        Draw bounding boxes and labels on the snapshot image.
        
        Args:
            img: PIL Image to annotate
            findings: Findings for this page
            page: fitz Page object (for coordinate reference)
            pix: Rendered pixmap (for coordinate scaling)
            
        Returns:
            Annotated image
        """
        draw = ImageDraw.Draw(img)
        
        # Try to load a font, fall back to default if not available
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 14)
            small_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 10)
        except:
            logger.warning("Could not load TrueType font, using default")
            font = ImageFont.load_default()
            small_font = ImageFont.load_default()
        
        # Color coding by status
        colors = {
            StatusCode.FAIL: (255, 0, 0),      # Red - Critical issues
            StatusCode.ATTN: (255, 165, 0),    # Orange - Warnings
            StatusCode.TBD: (255, 255, 0),     # Yellow - Needs verification
            StatusCode.OK: (0, 255, 0),        # Green - Passed
            StatusCode.FYI: (128, 128, 128),   # Gray - Informational
        }
        
        issue_num = 1
        
        for finding in findings:
            if not finding.matched_runs:
                continue
            
            # Get color based on status
            color = colors.get(finding.status, (128, 128, 128))
            
            # Draw each matched run's bounding box
            for run in finding.matched_runs:
                if not run.bbox:
                    continue
                
                # Scale bbox from PDF coordinates to image coordinates
                scale_x = pix.width / page.rect.width
                scale_y = pix.height / page.rect.height
                
                x0, y0, x1, y1 = run.bbox
                x0_img = int(x0 * scale_x)
                y0_img = int(y0 * scale_y)
                x1_img = int(x1 * scale_x)
                y1_img = int(y1 * scale_y)
                
                # Draw bounding box rectangle
                draw.rectangle(
                    [(x0_img, y0_img), (x1_img, y1_img)],
                    outline=color,
                    width=3
                )
                
                # Draw issue number label with background
                label = f"#{issue_num}"
                
                # Get text bounding box for background
                try:
                    bbox = draw.textbbox((x0_img, y0_img - 20), label, font=font)
                    # Draw background rectangle
                    draw.rectangle(bbox, fill=color)
                    # Draw text
                    draw.text(
                        (x0_img, y0_img - 20),
                        label,
                        fill=(255, 255, 255),  # White text
                        font=font
                    )
                except:
                    # Fallback if textbbox not available
                    draw.text(
                        (x0_img, y0_img - 15),
                        label,
                        fill=color,
                        font=font
                    )
            
            issue_num += 1
        
        return img

# =============================================================================
# MARKDOWN RENDERER
# =============================================================================

class MarkdownRenderer:
    """
    Renders verification results to markdown format.
    
    Generates complete report with all sections formatted
    as markdown tables for display in ChatGPT.
    """
    
    @staticmethod
    def render_complete_report(
        copy_doc: CopyDocument,
        artwork_extraction: ExtractionResult,
        copy_quality_issues: List[CopyQualityIssue],
        claim_risks: List[ClaimRisk],
        conversions: List[ConversionCheck],
        match_findings: List[MatchFinding],
        font_measurements: List[FontMeasurement],
        barcode_results: List[BarcodeResult],
        snapshot_paths: List[str],
        project_name: str = "Artwork Check"
    ) -> str:
        """
        Render complete verification report.
        
        Args:
            copy_doc: Parsed copy document
            artwork_extraction: Artwork extraction result
            copy_quality_issues: Section 3A issues
            claim_risks: Section 3B risks
            conversions: Section 3C conversions
            match_findings: Section 3D findings
            font_measurements: Section 3E fonts
            barcode_results: Section 3F barcodes
            snapshot_paths: Section 3G snapshot files
            project_name: Name of project
            
        Returns:
            Complete markdown report string
        """
        sections = []
        
        # Header
        sections.append("# Artwork Verification Report\n")
        sections.append(f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
        sections.append(f"*Checker Version: {config.VERSION}*\n")
        
        # Section 1: Project Header
        sections.append(MarkdownRenderer._render_section1(project_name, copy_doc))
        
        # Section 2: Files
        sections.append(MarkdownRenderer._render_section2(copy_doc, artwork_extraction))
        
        # Section 3: Core Verification Tables
        sections.append("\n## 3️⃣ Core Verification Tables\n")
        
        # 3A: Copy Quality
        sections.append(MarkdownRenderer._render_section3a(copy_quality_issues))
        
        # 3B: Claim Risk
        sections.append(MarkdownRenderer._render_section3b(claim_risks))
        
        # 3C: Conversion
        sections.append(MarkdownRenderer._render_section3c(conversions))
        
        # 3D: Artwork Match
        sections.append(MarkdownRenderer._render_section3d(match_findings))
        
        # 3E: Font Size
        sections.append(MarkdownRenderer._render_section3e(font_measurements))
        
        # 3F: Barcode
        sections.append(MarkdownRenderer._render_section3f(barcode_results))
        
        # 3G: Visual Snapshots
        sections.append(MarkdownRenderer._render_section3g(snapshot_paths, match_findings))
        
        # 3H: Score & Summary
        sections.append(MarkdownRenderer._render_section3h(match_findings, copy_quality_issues))
        
        # Section 4: Optional Fields
        sections.append(MarkdownRenderer._render_section4())
        
        # Section 5: Special Notes
        sections.append(MarkdownRenderer._render_section5(match_findings, artwork_extraction))
        
        return "\n".join(sections)
    
    @staticmethod
    def _render_section1(project_name: str, copy_doc: CopyDocument) -> str:
        """Render Section 1: Project Header."""
        return f"""
## 1️⃣ Project Header

| Field | Value |
|-------|-------|
| Project Name | {sanitize_for_markdown(project_name)} |
| Round / Version | {datetime.now().strftime('%m.%d.%y')} |
| Copy Document | {Path(copy_doc.file_path).name} |
| Languages | {', '.join(copy_doc.metadata.get('languages_found', ['EN']))} |
"""
    
    @staticmethod
    def _render_section2(copy_doc: CopyDocument, artwork: ExtractionResult) -> str:
        """Render Section 2: Files."""
        copy_status = "✅" if not copy_doc.warnings else "⚠️"
        art_status = "✅" if artwork.confidence > 0.5 else "⚠️"
        
        art_note = ""
        if artwork.extraction_method == ExtractionMethod.FAILED:
            art_note = "Text extraction failed - visual verification required"
        elif artwork.confidence < 1.0:
            art_note = f"Extraction confidence: {artwork.confidence:.0%}"
        
        return f"""
## 2️⃣ Files

| Type | Filename | Status | Note |
|------|----------|--------|------|
| Copy Document | {Path(copy_doc.file_path).name} | {copy_status} | {len(copy_doc.fields)} fields extracted |
| Artwork | {Path(artwork.file_path).name} | {art_status} | {sanitize_for_markdown(art_note)} |
"""
    
    @staticmethod
    def _render_section3a(issues: List[CopyQualityIssue]) -> str:
        """Render Section 3A: Copy Quality."""
        lines = ["\n### A. Copy Quality\n"]
        lines.append("| Language | Field | Issue Type | Recommendation | Status |")
        lines.append("|----------|-------|------------|----------------|--------|")
        
        if not issues:
            lines.append("| EN | All | — | No copy quality issues detected | ✅ |")
        else:
            for issue in issues:
                status_emoji = config.STATUS_EMOJI.get(issue.status_code.value, "❓")
                lines.append(
                    f"| {issue.language} | "
                    f"{sanitize_for_markdown(truncate_text(issue.field_name, 30))} | "
                    f"{issue.issue_type} | "
                    f"{sanitize_for_markdown(truncate_text(issue.recommendation, 50))} | "
                    f"{status_emoji} |"
                )
        
        return "\n".join(lines)
    
    @staticmethod
    def _render_section3b(risks: List[ClaimRisk]) -> str:
        """Render Section 3B: Claim Risk."""
        lines = ["\n### B. Claim Risk\n"]
        lines.append("| Language | Claim | Risk Level | Rationale | Regions | Action | Status |")
        lines.append("|----------|-------|------------|-----------|---------|--------|--------|")
        
        if not risks:
            lines.append("| — | No claims detected | — | — | — | — | ✅ |")
        else:
            for risk in risks:
                status_emoji = config.STATUS_EMOJI.get(risk.status_code.value, "❓")
                lines.append(
                    f"| {risk.language} | "
                    f"{sanitize_for_markdown(truncate_text(risk.claim_text, 40))} | "
                    f"{risk.risk_level.value} | "
                    f"{sanitize_for_markdown(truncate_text(risk.rationale, 40))} | "
                    f"{', '.join(risk.regions[:3])} | "
                    f"{risk.recommended_action} | "
                    f"{status_emoji} |"
                )
        
        return "\n".join(lines)
    
    @staticmethod
    def _render_section3c(conversions: List[ConversionCheck]) -> str:
        """Render Section 3C: Conversion."""
        lines = ["\n### C. Label-Claim Conversion\n"]
        lines.append("| Source | Declared (mL) | Calculated (fl oz) | Declared (fl oz) | Within ±0.10 | Status | Notes |")
        lines.append("|--------|---------------|-------------------|------------------|--------------|--------|-------|")
        
        if not conversions:
            lines.append("| — | — | — | — | — | ✅ | No conversions to check |")
        else:
            for conv in conversions:
                status_emoji = config.STATUS_EMOJI.get(conv.status_code.value, "❓")
                within = "Yes" if conv.within_tolerance else "No"
                lines.append(
                    f"| {conv.source_field} | "
                    f"{conv.declared_ml} | "
                    f"{conv.calculated_floz} | "
                    f"{conv.declared_floz} | "
                    f"{within} | "
                    f"{status_emoji} | "
                    f"{conv.notes} |"
                )
        
        return "\n".join(lines)
    
    @staticmethod
    def _render_section3d(findings: List[MatchFinding]) -> str:
        """Render Section 3D: Artwork Match."""
        lines = ["\n### D. Artwork Match\n"]
        
        # Group by panel and language
        grouped = {}
        for f in findings:
            key = (f.panel, f.language)
            if key not in grouped:
                grouped[key] = []
            grouped[key].append(f)
        
        for (panel, language), panel_findings in grouped.items():
            lines.append(f"\n**{panel} – {language}**\n")
            lines.append("| Field | Copy Doc Value | Artwork Value | Match | Notes |")
            lines.append("|-------|----------------|---------------|-------|-------|")
            
            for finding in panel_findings:
                # Determine match emoji
                if finding.match_type == MatchType.EXACT_MATCH:
                    match_emoji = "✅"
                elif finding.match_type == MatchType.NEAR_MATCH:
                    match_emoji = "⚠️"
                else:
                    match_emoji = "❌"
                
                # Format artwork value
                art_val = finding.artwork_value or "NOT FOUND"
                
                # Add issue ID to notes if present
                notes = "; ".join(finding.notes) if finding.notes else ""
                if finding.issue_id:
                    notes = f"[{finding.issue_id}] {notes}"
                
                lines.append(
                    f"| {sanitize_for_markdown(truncate_text(finding.field_name, 25))} | "
                    f"{sanitize_for_markdown(truncate_text(finding.copy_value, 40))} | "
                    f"{sanitize_for_markdown(truncate_text(art_val, 40))} | "
                    f"{match_emoji} | "
                    f"{sanitize_for_markdown(truncate_text(notes, 60))} |"
                )
        
        return "\n".join(lines)
    
    @staticmethod
    def _render_section3e(measurements: List[FontMeasurement]) -> str:
        """Render Section 3E: Font Size."""
        lines = ["\n### E. Font Size\n"]
        lines.append("| Element | Size (pt) | Requirement | Status |")
        lines.append("|---------|-----------|-------------|--------|")
        
        if not measurements:
            lines.append("| — | — | — | 🔍 Unable to extract font metadata |")
        else:
            # Find smallest font
            smallest = min(measurements, key=lambda m: m.font_size_pt)
            
            # Check against requirements
            usa_min = 4.5
            eu_min = 6.0
            
            usa_status = "✅" if smallest.font_size_pt >= usa_min else "❌"
            eu_status = "✅" if smallest.font_size_pt >= eu_min else "❌"
            
            lines.append(
                f"| Smallest text: \"{sanitize_for_markdown(truncate_text(smallest.text, 20))}\" | "
                f"{smallest.font_size_pt} | "
                f"USA ≥{usa_min}pt | {usa_status} |"
            )
            lines.append(
                f"| (same) | {smallest.font_size_pt} | EU ≥{eu_min}pt | {eu_status} |"
            )
        
        return "\n".join(lines)
    
    @staticmethod
    def _render_section3f(results: List[BarcodeResult]) -> str:
        """Render Section 3F: Barcode."""
        lines = ["\n### F. Barcode\n"]
        lines.append("| Symbology | Decoded Digits | Check Digit | Scan Status | Notes |")
        lines.append("|-----------|----------------|-------------|-------------|-------|")
        
        if not results:
            lines.append("| N/A | — | — | — | No barcode scanned |")
        else:
            for result in results:
                status_emoji = config.STATUS_EMOJI.get(result.status_code.value, "❓")
                check = "✅ Valid" if result.check_digit_valid else "❌ Invalid"
                if not result.scan_successful:
                    check = "—"
                
                lines.append(
                    f"| {result.symbology} | "
                    f"{result.decoded_digits or '—'} | "
                    f"{check} | "
                    f"{'Scanned' if result.scan_successful else 'Failed'} | "
                    f"{sanitize_for_markdown(result.quality_notes)} |"
                )
        
        return "\n".join(lines)
    
    @staticmethod
    def _render_section3g(snapshot_paths: List[str], findings: List[MatchFinding]) -> str:
        """Render Section 3G: Visual Snapshots."""
        lines = ["\n### G. Visual Snapshots\n"]
        
        # Count issues that should have snapshots
        issues_with_bbox = [f for f in findings if f.status_code != StatusCode.OK and f.bbox]
        
        if not snapshot_paths:
            if issues_with_bbox:
                lines.append(f"*{len(issues_with_bbox)} issues flagged for visual verification.*\n")
                lines.append("*Snapshots not generated in this run. Please visually verify flagged items.*\n")
            else:
                lines.append("*No issues require visual snapshots.*\n")
        else:
            lines.append(f"*Generated {len(snapshot_paths)} annotated snapshot(s).*\n")
            for path in snapshot_paths:
                lines.append(f"- `{Path(path).name}`")
        
        # Reference table
        lines.append("\n| ID | Issue | Location |")
        lines.append("|----|-------|----------|")
        
        for finding in findings:
            if finding.issue_id:
                lines.append(
                    f"| {finding.issue_id} | "
                    f"{finding.match_type.value} | "
                    f"{finding.panel} |"
                )
        
        return "\n".join(lines)
    
    @staticmethod
    def _render_section3h(findings: List[MatchFinding], quality_issues: List[CopyQualityIssue]) -> str:
        """Render Section 3H: Score & Summary."""
        total_checks = len(findings)
        exact_matches = sum(1 for f in findings if f.match_type == MatchType.EXACT_MATCH)
        near_matches = sum(1 for f in findings if f.match_type == MatchType.NEAR_MATCH)
        fails = sum(1 for f in findings if f.status_code == StatusCode.FAIL)
        
        score = (exact_matches / total_checks * 100) if total_checks > 0 else 0
        
        lines = ["\n### H. Score & Summary\n"]
        lines.append("| Area | Checks | Matches | Score % | Notes |")
        lines.append("|------|--------|---------|---------|-------|")
        
        lines.append(
            f"| Artwork Match | {total_checks} | {exact_matches} | {score:.1f}% | "
            f"{near_matches} near, {fails} failed |"
        )
        
        quality_ok = len([i for i in quality_issues if i.status_code == StatusCode.OK])
        quality_total = len(quality_issues) if quality_issues else 1
        quality_score = (quality_ok / quality_total * 100) if quality_total > 0 else 100
        
        lines.append(
            f"| Copy Quality | {quality_total} | {quality_ok} | {quality_score:.1f}% | — |"
        )
        
        overall = (score + quality_score) / 2
        lines.append(f"| **Overall** | — | — | **{overall:.1f}%** | — |")
        
        # Top fixes
        lines.append("\n**Top Fixes (❌):**\n")
        fail_findings = [f for f in findings if f.status_code == StatusCode.FAIL][:5]
        if fail_findings:
            for f in fail_findings:
                lines.append(f"- [{f.issue_id}] {f.field_name}: {f.match_type.value}")
        else:
            lines.append("- None")
        
        # Attention items
        lines.append("\n**Attention (⚠️):**\n")
        attn_findings = [f for f in findings if f.status_code == StatusCode.ATTN][:5]
        if attn_findings:
            for f in attn_findings:
                lines.append(f"- {f.field_name}: {'; '.join(f.notes[:1])}")
        else:
            lines.append("- None")
        
        return "\n".join(lines)
    
    @staticmethod
    def _render_section4() -> str:
        """Render Section 4: Optional Fields."""
        return """
## 4️⃣ Optional Fields

| Field | Content |
|-------|---------|
| Version Change Log | Not provided |
| Creative Brand-Voice Check | Not requested |
| One-Page PDF Summary Export | Available on request |
"""
    
    @staticmethod
    def _render_section5(findings: List[MatchFinding], artwork: ExtractionResult) -> str:
        """Render Section 5: Special Notes."""
        lines = ["\n## 5️⃣ Special Notes / Constraints\n"]
        lines.append("| Constraint | Source | Applies To | Notes |")
        lines.append("|------------|--------|------------|-------|")
        
        lines.append(
            "| Text must match character-for-character | Brand | All panels | "
            "Including punctuation, case, diacritics |"
        )
        
        # Add note if extraction had issues
        if artwork.extraction_method == ExtractionMethod.FAILED:
            lines.append(
                "| Visual verification required | Technical | All fields | "
                "Text extraction failed (outlined fonts likely) |"
            )
        
        # Count visual verification items
        verify_count = sum(1 for f in findings if f.requires_zoom)
        if verify_count > 0:
            lines.append(
                f"| {verify_count} fields require visual verification | Zoom Triggers | "
                f"See flagged items | Numbers, units, or low confidence |"
            )
        
        return "\n".join(lines)




# =============================================================================
# MAIN ORCHESTRATOR
# =============================================================================

class ArtworkChecker:
    """
    Main orchestrator for the artwork verification process.
    
    Coordinates all extraction, analysis, and reporting
    components to produce a comprehensive verification report.
    """
    
    def __init__(self):
        """Initialize all component extractors and analyzers."""
        self.copy_extractor = CopyDocumentExtractor()
        self.pdf_extractor = PDFExtractor()
        self.ai_extractor = AIExtractor()
        self.font_extractor = FontExtractor()
        self.matcher = ArtworkMatcher(Config)
    
    def run_check(
        self,
        copy_path: Path,
        artwork_path: Path,
        output_dir: Optional[Path] = None,
        project_name: Optional[str] = None
    ) -> str:
        """
        Run complete artwork verification check.
        
        Args:
            copy_path: Path to copy document (.docx)
            artwork_path: Path to artwork file (.pdf or .ai)
            output_dir: Optional directory for output files
            project_name: Optional project name for report
            
        Returns:
            Complete markdown report
            
        Raises:
            FileNotFoundError: If input files don't exist
            ValueError: If file formats are invalid
        """
        logger.info("=" * 60)
        logger.info("ARTWORK CHECKER v2.0.0")
        logger.info("=" * 60)
        
        # Validate inputs
        if not copy_path.exists():
            raise FileNotFoundError(f"Copy document not found: {copy_path}")
        if not artwork_path.exists():
            raise FileNotFoundError(f"Artwork file not found: {artwork_path}")
        
        # Set defaults
        if output_dir is None:
            output_dir = Path("./output")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if project_name is None:
            project_name = artwork_path.stem
        
        # =====================================================================
        # STEP 1: Extract Copy Document
        # =====================================================================
        logger.info("\n[1/7] Extracting copy document...")
        copy_doc = self.copy_extractor.extract(copy_path)
        
        logger.info(f"  - Fields extracted: {len(copy_doc.fields)}")
        logger.info(f"  - Languages: {copy_doc.metadata.get('languages_found', [])}")
        logger.info(f"  - Strikethrough items: {len(copy_doc.strikethrough_text)}")
        logger.info(f"  - Legacy items: {len(copy_doc.legacy_text)}")
        logger.info(f"  - Instructional notes: {len(copy_doc.instructional_notes)}")
        
        if copy_doc.warnings:
            for warning in copy_doc.warnings:
                logger.warning(f"  ⚠️ {warning}")
        
        # =====================================================================
        # STEP 2: Extract Artwork
        # =====================================================================
        logger.info("\n[2/7] Extracting artwork...")
        
        artwork_ext = artwork_path.suffix.lower()
        
        if artwork_ext == '.pdf':
            artwork_extraction = self.pdf_extractor.extract(artwork_path)
        elif artwork_ext == '.ai':
            artwork_extraction = self.ai_extractor.extract(artwork_path)
        else:
            raise ValueError(f"Unsupported artwork format: {artwork_ext}")
        
        logger.info(f"  - Method: {artwork_extraction.extraction_method.value}")
        logger.info(f"  - Confidence: {artwork_extraction.confidence:.0%}")
        logger.info(f"  - Text runs: {len(artwork_extraction.text_runs)}")
        logger.info(f"  - Pages: {artwork_extraction.pages_processed}")
        
        if artwork_extraction.warnings:
            for warning in artwork_extraction.warnings[:3]:
                logger.warning(f"  ⚠️ {warning}")
        
        # =====================================================================
        # STEP 3: Copy Quality Analysis (Section 3A)
        # =====================================================================
        logger.info("\n[3/7] Analyzing copy quality...")
        copy_quality_issues = CopyQualityChecker.analyze_copy_document(copy_doc)
        
        fails = sum(1 for i in copy_quality_issues if i.status_code == StatusCode.FAIL)
        attns = sum(1 for i in copy_quality_issues if i.status_code == StatusCode.ATTN)
        logger.info(f"  - Issues found: {len(copy_quality_issues)} ({fails} fails, {attns} warnings)")
        
        # =====================================================================
        # STEP 4: Claim Risk Assessment (Section 3B)
        # =====================================================================
        logger.info("\n[4/7] Assessing claim risks...")
        claim_risks = ClaimRiskAssessor.assess_all_claims(copy_doc)
        
        high_risk = sum(1 for r in claim_risks if r.risk_level == RiskLevel.HIGH)
        logger.info(f"  - Claims assessed: {len(claim_risks)} ({high_risk} high-risk)")
        
        # =====================================================================
        # STEP 5: Match Artwork Against Copy (Sections 3C, 3D)
        # =====================================================================
        logger.info("\n[5/7] Matching artwork against copy document...")
        
        # Unit conversions (Section 3C)
        conversions = ConversionChecker.check_conversions(copy_doc)
        logger.info(f"  - Conversions checked: {len(conversions)}")
        
        # Artwork matching (Section 3D)
        match_findings = self.matcher.match_all_fields(
            copy_doc.fields,
            artwork_extraction
        )
        
        exact = sum(1 for f in match_findings if f.match_type == MatchType.EXACT_MATCH)
        near = sum(1 for f in match_findings if f.match_type == MatchType.NEAR_MATCH)
        missing = sum(1 for f in match_findings if f.match_type == MatchType.MISSING_IN_ARTWORK)
        
        logger.info(f"  - Exact matches: {exact}")
        logger.info(f"  - Near matches: {near}")
        logger.info(f"  - Missing: {missing}")
        
        # =====================================================================
        # STEP 6: Extract Font & Barcode Data (Sections 3E, 3F)
        # =====================================================================
        logger.info("\n[6/7] Extracting font and barcode data...")
        
        # Font measurements (Section 3E)
        font_measurements = []
        if artwork_ext == '.pdf':
            font_measurements = self.font_extractor.extract_from_pdf(artwork_path)
            if font_measurements:
                smallest = self.font_extractor.get_smallest_font(font_measurements)
                logger.info(f"  - Smallest font: {smallest.font_size_pt}pt")
        
        # Barcode scanning (Section 3F)
        barcode_results = []
        if artwork_ext == '.pdf':
            barcode_results = BarcodeScanner.scan_from_pdf(artwork_path)
            scanned = sum(1 for b in barcode_results if b.scan_successful)
            logger.info(f"  - Barcodes found: {scanned}")
        
        # =====================================================================
        # STEP 7: Generate Snapshots (Section 3G)
        # =====================================================================
        logger.info("\n[7/7] Generating visual snapshots...")
        
        snapshot_paths = []
        if config.SNAPSHOT_ENABLED and artwork_ext == '.pdf':
            snapshot_paths = SnapshotGenerator.generate_snapshots_for_findings(
                artwork_path,
                match_findings,
                output_dir / "snapshots"
            )
            logger.info(f"  - Snapshots generated: {len(snapshot_paths)}")
        
        # =====================================================================
        # GENERATE REPORT
        # =====================================================================
        logger.info("\n" + "=" * 60)
        logger.info("GENERATING REPORT")
        logger.info("=" * 60)
        
        report = MarkdownRenderer.render_complete_report(
            copy_doc=copy_doc,
            artwork_extraction=artwork_extraction,
            copy_quality_issues=copy_quality_issues,
            claim_risks=claim_risks,
            conversions=conversions,
            match_findings=match_findings,
            font_measurements=font_measurements,
            barcode_results=barcode_results,
            snapshot_paths=snapshot_paths,
            project_name=project_name
        )
        
        # Save report
        report_path = output_dir / "artwork_check_report.md"
        report_path.write_text(report, encoding='utf-8')
        logger.info(f"Report saved: {report_path}")
        
        # =====================================================================
        # SUMMARY
        # =====================================================================
        total_checks = len(match_findings)
        accuracy = (exact / total_checks * 100) if total_checks > 0 else 0
        
        logger.info("\n" + "=" * 60)
        logger.info("SUMMARY")
        logger.info("=" * 60)
        logger.info(f"Total fields checked: {total_checks}")
        logger.info(f"Exact matches: {exact} ✅")
        logger.info(f"Near matches: {near} ⚠️")
        logger.info(f"Missing/Failed: {missing + (total_checks - exact - near - missing)} ❌")
        logger.info(f"Accuracy: {accuracy:.1f}%")
        logger.info("=" * 60)
        
        return report


# =============================================================================
# COMMAND LINE INTERFACE
# =============================================================================

def parse_arguments() -> argparse.Namespace:
    """
    Parse command line arguments.
    
    Returns:
        Parsed argument namespace
    """
    parser = argparse.ArgumentParser(
        description="Artwork Checker v2.0.0 - Verify artwork against copy documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python artwork_checker_v2.py --copy Body_Butter_Copy_Doc.docx --artwork Carton_AW.pdf
    python artwork_checker_v2.py --copy copy.docx --artwork artwork.ai --output ./reports
    python artwork_checker_v2.py --copy copy.docx --artwork art.pdf --project "Big Embrace"
        """
    )
    
    parser.add_argument(
        '--copy', '-c',
        type=Path,
        required=True,
        help='Path to copy document (.docx file)'
    )
    
    parser.add_argument(
        '--artwork', '-a',
        type=Path,
        required=True,
        help='Path to artwork file (.pdf or .ai file)'
    )
    
    parser.add_argument(
        '--output', '-o',
        type=Path,
        default=Path('./output'),
        help='Output directory for reports (default: ./output)'
    )
    
    parser.add_argument(
        '--project', '-p',
        type=str,
        default=None,
        help='Project name for report header'
    )
    
    parser.add_argument(
        '--verbose', '-v',
        action='store_true',
        help='Enable verbose logging'
    )
    
    parser.add_argument(
        '--no-snapshots',
        action='store_true',
        help='Disable snapshot generation'
    )
    
    parser.add_argument(
        '--version',
        action='version',
        version=f'Artwork Checker v{config.VERSION}'
    )
    
    return parser.parse_args()


def main():
    """
    Main entry point for command line execution.
    
    Parses arguments, runs the artwork check, and outputs results.
    """
    args = parse_arguments()
    
    # Configure logging
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Configure snapshots
    if args.no_snapshots:
        config.SNAPSHOT_ENABLED = False
    
    # Run checker
    try:
        checker = ArtworkChecker()
        report = checker.run_check(
            copy_path=args.copy,
            artwork_path=args.artwork,
            output_dir=args.output,
            project_name=args.project
        )
        
        # Print report to stdout
        print("\n" + "=" * 60)
        print("MARKDOWN REPORT")
        print("=" * 60 + "\n")
        print(report)
        
        sys.exit(0)
        
    except FileNotFoundError as e:
        logger.error(f"File not found: {e}")
        sys.exit(1)
        
    except ValueError as e:
        logger.error(f"Invalid input: {e}")
        sys.exit(1)
        
    except Exception as e:
        logger.error(f"Unexpected error: {e}", exc_info=True)
        sys.exit(2)


# =============================================================================
# ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    main()
